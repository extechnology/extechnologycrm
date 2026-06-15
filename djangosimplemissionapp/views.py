from urllib import request

from django.contrib.auth.models import Permission
from rest_framework import views, viewsets, filters, serializers
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.generics import ListCreateAPIView, RetrieveUpdateDestroyAPIView, ListAPIView
from rest_framework.views import APIView
from rest_framework.exceptions import PermissionDenied
from django.db.models import Q, Subquery, OuterRef
from .models import (
    User,ProjectClient,ProjectBusinessAddress,DomainOrServerThirdPartyServiceProvider,Role,
    ProjectDomain,ProjectServer,ProjectFinance,Team,ProjectTeam,ProjectNature,
    Project,ProjectBaseInformation,ProjectExcution,ProjectTeamMember,ProjectService,ProjectServiceMember,
    EmployeeDailyActivity,ActivityLog,Invoice,InvoiceItem,Payment,ActivityExceedComment,
    Notification,EmployeeLeave,Company,CompanyProfile,Salary,Attendance,Employee,OtherIncome,OtherExpense,ProjectDocument,
    ClientAdvance, UserSalary, SalaryIncrement,ProjectExbot,Lead,FollowUp,Schedule,SystemAuditLog,LoginUserDetails,Device
)
  
from .serializers import (
    UserSerializer, ChangePasswordSerializer, AdminChangePasswordSerializer,ProjectClientSerializer,
    ProjectBusinessAddressSerializer,DomainOrServerThirdPartyServiceProviderSerializer,
    ProjectDomainSerializer,ProjectServerSerializer,ProjectFinanceSerializer,TeamSerializer,
    ProjectTeamSerializer,ProjectNatureSerializer,
    ProjectSerializer,ProjectBaseInformationSerializer,ProjectExcutionSerializer,ProjectTeamMemberSerializer,ProjectServiceSerializer,
    EmployeeDailyActivitySerializer,ActivityLogSerializer,InvoiceSerializer,InvoiceItemSerializer,PaymentSerializer,ActivityExceedCommentSerializer,
    NotificationSerializer,EmployeeLeaveSerializer,CompanySerializer,CompanyProfileSerializer,SalarySerializer,AttendanceSerializer,EmployeeSerializer,OtherIncomeSerializer,OtherExpenseSerializer,RoleSerializer, PermissionSerializer,
    ProjectDocumentSerializer, ProjectSummarySerializer, ClientAdvanceSerializer, ClientSummarySerializer, UserSalarySerializer, SalaryIncrementSerializer,ProjectExbotSerializer,
    LeadSerializer, FollowUpSerializer, ScheduleSerializer, EmployeeSalarySummarySerializer, LoginUserDetailsSerializer, DeviceSerializer
)

from rest_framework.decorators import action
from .utils import get_date_filter_q
from .pdf_utils import generate_invoice_pdf
from django.contrib.auth import authenticate
from rest_framework import status
from django.http import Http404, JsonResponse, FileResponse
from .permissions import IsSuperAdmin, IsDeveloper, IsAdmin
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework_simplejwt.serializers import TokenObtainPairSerializer
from django.utils import timezone
from decimal import Decimal
from datetime import timedelta
from rest_framework.pagination import PageNumberPagination
from rest_framework.views import APIView

class StandardResultsSetPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100

class UserListAPIView(views.APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Admins or users with specific permission can see all users
        is_privileged = any(role.upper() in ['SUPERADMIN', 'ADMIN', 'BILLING'] for role in request.user.role_names) or \
                        request.user.has_perm('djangosimplemissionapp.viewall_user') 
        
        
        if is_privileged:
            users = User.objects.all()
        else:
            # Other users can only see themselves
            users = User.objects.filter(id=request.user.id)
        
        search_query = request.query_params.get('search', None)
        if search_query:
            # Check if search query is like USR-0001
            id_filter = Q()
            if search_query.upper().startswith('USR-'):
                try:
                    numeric_id = int(search_query.split('-')[1])
                    id_filter = Q(id=numeric_id)
                except (IndexError, ValueError):
                    pass
            elif search_query.isdigit():
                id_filter = Q(id=int(search_query))

            users = users.filter(
                id_filter |
                Q(username__icontains=search_query) |
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query) |
                Q(email__icontains=search_query) |
                Q(designation__icontains=search_query) |
                Q(phone_number__icontains=search_query)
            )
            
        # Additional filters
        designation = request.query_params.get('designation')
        if designation:
            users = users.filter(designation__icontains=designation)
            
        role = request.query_params.get('role')
        if role:
            users = users.filter(role__name__icontains=role)
            
        status_param = request.query_params.get('status')
        if status_param:
            if status_param.lower() == 'active':
                users = users.filter(is_active=True)
            elif status_param.lower() == 'inactive':
                users = users.filter(is_active=False)
                
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if start_date and end_date:
            users = users.filter(date_joined__date__range=[start_date, end_date])
        elif start_date:
            users = users.filter(date_joined__date__gte=start_date)
        elif end_date:
            users = users.filter(date_joined__date__lte=end_date)
            
        users = users.order_by('-date_joined')
            
        paginator = StandardResultsSetPagination()
        result_page = paginator.paginate_queryset(users, request)
        serializer = UserSerializer(result_page, many=True, context={'request': request})
        return paginator.get_paginated_response(serializer.data)

    def post(self, request):
        # Only SuperAdmin can create users
        if not request.user.has_role('SuperAdmin'):
             return Response({'error': 'Permission denied. Only SuperAdmin can create users.'}, status=status.HTTP_403_FORBIDDEN)

        serializer = UserSerializer(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class UserDetailAPIView(views.APIView):
    permission_classes = [IsAuthenticated]  

    def get_object(self, pk):
        try:
            return User.objects.get(pk=pk)
        except User.DoesNotExist:
            raise Http404

    def get(self, request, pk):
        user = self.get_object(pk)
        
        # Privacy check: Only SuperAdmin/Admin or users with specific permission can view other profiles
        is_privileged = any(role.upper() in ['SUPERADMIN', 'ADMIN', 'BILLING'] for role in request.user.role_names) or \
                        request.user.has_perm('djangosimplemissionapp.view_all_employee_performance') or \
                        request.user.has_perm('djangosimplemissionapp.view_all_activities') or \
                        request.user.has_perm('djangosimplemissionapp.view_all_team_performance') or \
                        request.user.has_perm('djangosimplemissionapp.viewall_user')
        
        if not is_privileged and request.user.id != user.id:
             return Response({'error': 'Permission denied. You can only view your own profile.'}, status=status.HTTP_403_FORBIDDEN)

        serializer = UserSerializer(user, context={'request': request})
        return Response(serializer.data)

    def put(self, request, pk):
        # Only SuperAdmin can edit users
        if not request.user.has_role('SuperAdmin'):
            return Response({'error': 'Permission denied. Only SuperAdmin can edit users.'}, status=status.HTTP_403_FORBIDDEN)
            
        user = self.get_object(pk)
        serializer = UserSerializer(user, data=request.data, partial=True, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        # Only SuperAdmin can delete users
        if not request.user.has_role('SuperAdmin'):
            return Response({'error': 'Permission denied. Only SuperAdmin can delete users.'}, status=status.HTTP_403_FORBIDDEN)

        user = self.get_object(pk)
        
        SystemAuditLog.objects.create(
            action="USER_DELETED",
            performed_by=f"{request.user.username} (ID: {request.user.id})",
            target=f"{user.username} (ID: {user.id})"
        )
        
        user.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class CurrentUserView(views.APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        serializer = UserSerializer(request.user, context={'request': request})
        return Response(serializer.data)

    def put(self, request):
        # Exclude role_id and other sensitive fields from self-update
        data = request.data.copy()
        data.pop('role_id', None)
        data.pop('is_superuser', None)
        data.pop('is_staff', None)
        
        serializer = UserSerializer(request.user, data=data, partial=True, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class CustomTokenObtainPairSerializer(TokenObtainPairSerializer):
    device_id = serializers.CharField(required=False, allow_blank=True)
    device_name = serializers.CharField(required=False, allow_blank=True)
    device_type = serializers.CharField(required=False, allow_blank=True)
    device_info = serializers.JSONField(required=False, allow_null=True)
    
    @classmethod
    def get_token(cls, user):
        token = super().get_token(user)
        token['token_version'] = user.token_version
        return token

    def validate(self, attrs):
        data = super().validate(attrs)
        user = self.user
        
        # Device-lock validation
        device_id = attrs.get('device_id') or self.context.get('request').data.get('device_id')
        device_name = attrs.get('device_name') or self.context.get('request').data.get('device_name')
        device_type = attrs.get('device_type') or self.context.get('request').data.get('device_type', 'other')
        device_info = attrs.get('device_info') or self.context.get('request').data.get('device_info', {})
        
        # Check if user has any approved devices
        has_approved_devices = user.devices.filter(is_approved=True).exists()
        
        if device_id:
            from .models import Device
            from django.utils import timezone
            
            try:
                device = Device.objects.get(device_id=device_id, user=user)
                
                # Check if device is approved
                if not device.is_approved:
                    raise serializers.ValidationError(
                        'admin permission required for signup',
                        code='device_not_approved'
                    )
                
                # Update device login info
                device.last_login = timezone.now()
                device.login_count += 1
                device.save(update_fields=['last_login', 'login_count'])
                
            except Device.DoesNotExist:
                # Device doesn't exist - create it (pending approval)
                if has_approved_devices:
                    # User has approved devices, so new device needs approval
                    Device.objects.create(
                        user=user,
                        device_id=device_id,
                        device_name=device_name or f"New {device_type.capitalize()} Device",
                        device_type=device_type,
                        device_info=device_info or {},
                        is_approved=False
                    )
                    raise serializers.ValidationError(
                        'admin permission required for signup',
                        code='device_pending_approval'
                    )  
                else:
                    # First device - auto-approve for now
                    from .models import Device
                    device = Device.objects.create(
                        user=user,
                        device_id=device_id,
                        device_name=device_name or f"Device {device_type}",
                        device_type=device_type,
                        device_info=device_info or {},
                        is_approved=True,  # Auto-approve first device
                        approved_by=None  # System auto-approval
                    )
                    device.last_login = timezone.now()
                    device.login_count = 1
                    device.save(update_fields=['last_login', 'login_count'])
        elif has_approved_devices:
            # User has approved devices but didn't provide device_id
            raise serializers.ValidationError(
                'Device ID is required. This account requires device approval.',
                code='device_id_required'
            )

        data['is_logged_in'] = True
        data['role'] = user.role.name if user.role else None
        data['permissions'] = list(user.role.permissions.values_list('codename', flat=True)) if user.role else []
        data['is_superuser'] = user.is_superuser

        data['user'] = {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'phone_number': user.phone_number,
            'designation': user.designation,
            'profile_pic': user.profile_pic.url if user.profile_pic else None,
            'is_superuser': user.is_superuser,
            'role': user.role.name if user.role else None,
            'roles': [{'name': user.role.name}] if user.role else [],
            'permissions': data['permissions']
        }
        return data


class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer
    
    def post(self, request, *args, **kwargs):
        """
        Override post to record login attempts in LoginUserDetails model.
        """
        from .utils import record_user_login
        
        response = super().post(request, *args, **kwargs)
        
        # Record successful login
        if response.status_code == 200:
            try:
                username = request.data.get('username')
                user = User.objects.get(username=username)
                record_user_login(user, request, login_status='SUCCESS')
            except Exception as e:
                # Log the error but don't fail the login
                print(f"Error recording login: {str(e)}")
        else:
            # Record failed login attempt
            try:
                username = request.data.get('username')
                user = User.objects.get(username=username)
                record_user_login(user, request, login_status='FAILED', notes='Invalid credentials')
            except:
                pass  # User doesn't exist, don't record
        
        return response

class ChangePasswordView(views.APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request):
        serializer = ChangePasswordSerializer(data=request.data)
        if serializer.is_valid():
            user = request.user
            if user.check_password(serializer.data.get('old_password')):
                user.set_password(serializer.data.get('new_password'))
                user.save()
                return Response({'status': 'password set'}, status=status.HTTP_200_OK)
            return Response({'old_password': ['Wrong password.']}, status=status.HTTP_400_BAD_REQUEST)


class LogoutView(views.APIView):
    """
    Logout view that records logout time and session duration in LoginUserDetails.
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        from .utils import record_user_logout
        
        try:
            # Record logout
            record_user_logout(request.user)
            
            return Response({
                'status': 'logout successful',
                'message': 'You have been logged out successfully'
            }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({
                'status': 'logout error',
                'message': f'Logout recorded with error: {str(e)}'
            }, status=status.HTTP_200_OK)  # Still return 200 even if there's an error


class DeviceListCreateAPIView(ListCreateAPIView):
    """
    List all devices for the current user or register a new device.
    POST: Register a new device (will be pending approval if user has other approved devices)
    GET: List all devices for the current user
    """
    serializer_class = DeviceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    
    def get_queryset(self):
        """Users can only see their own devices, admins can see all (with optional user filtering)"""
        user = self.request.user
        if user.is_superuser or user.has_perm('djangosimplemissionapp.view_all_devices'):
            queryset = Device.objects.all()
            user_id = self.request.query_params.get('user')
            if user_id:
                queryset = queryset.filter(user_id=user_id)
            return queryset
        return user.devices.all()

    def delete(self, request, *args, **kwargs):
        """Delete all device records. Restricted to superuser or users with delete_device permission."""
        user = request.user
        if not (user.is_superuser or user.has_perm('djangosimplemissionapp.delete_device')):
            return Response({'error': 'You do not have permission to delete all devices.'}, status=status.HTTP_403_FORBIDDEN)
        
        # Delete all devices
        deleted_count, _ = Device.objects.all().delete()
        return Response({'message': f'Successfully deleted all {deleted_count} devices.'}, status=status.HTTP_200_OK)
    
    def create(self, request, *args, **kwargs):
        """Register a new device"""
        device_id = request.data.get('device_id')
        device_name = request.data.get('device_name')
        device_type = request.data.get('device_type', 'other')
        device_info = request.data.get('device_info', {})
        
        if not device_id:
            return Response({'error': 'device_id is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Check if device already exists for this user
            device = Device.objects.get(device_id=device_id, user=request.user)
            return Response({
                'error': 'Device already registered',
                'device': DeviceSerializer(device).data
            }, status=status.HTTP_400_BAD_REQUEST)
        except Device.DoesNotExist:
            # Create new device
            has_approved_devices = request.user.devices.filter(is_approved=True).exists()
            
            # First device is auto-approved, subsequent devices need admin approval
            is_approved = not has_approved_devices
            
            device = Device.objects.create(
                user=request.user,
                device_id=device_id,
                device_name=device_name or f"Device {device_type}",
                device_type=device_type,
                device_info=device_info or {},
                is_approved=is_approved
            )
            
            serializer = self.get_serializer(device)
            return Response(
                {
                    'message': 'Device registered successfully' if not is_approved else 'First device auto-approved',
                    'device': serializer.data
                },
                status=status.HTTP_201_CREATED
            )


class DeviceDetailAPIView(RetrieveUpdateDestroyAPIView):
    """
    Retrieve, update, or delete a specific device.
    """
    serializer_class = DeviceSerializer
    permission_classes = [IsAuthenticated]
    lookup_field = 'id'
    
    def get_queryset(self):
        """Users can only access their own devices"""
        user = self.request.user
        if user.is_superuser or user.has_perm('djangosimplemissionapp.view_all_devices'):
            return Device.objects.all()
        return user.devices.all()
    
    def perform_destroy(self, instance):
        """Delete a device"""
        # Prevent deletion of all devices - user must have at least one
        if instance.user.devices.count() == 1:
            raise PermissionDenied('Cannot delete the only device. Register another device first.')
        instance.delete()


class DeviceApprovalAPIView(views.APIView):
    """
    Admin endpoint to approve or reject devices for users.
    POST with action='approve' or action='reject'
    """
    permission_classes = [IsAuthenticated]
    
    def post(self, request, device_id):
        # Check if user has permission to manage devices
        if not (request.user.is_superuser or request.user.has_perm('djangosimplemissionapp.manage_device_approvals')):
            return Response(
                {'error': 'You do not have permission to approve/reject devices'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            device = Device.objects.get(id=device_id)
        except Device.DoesNotExist:
            return Response({'error': 'Device not found'}, status=status.HTTP_404_NOT_FOUND)
        
        action = request.data.get('action')  # 'approve' or 'reject'
        reason = request.data.get('reason', '')
        
        if action == 'approve':
            device.approve(request.user, reason)
            return Response({
                'message': 'Device approved successfully',
                'device': DeviceSerializer(device).data
            }, status=status.HTTP_200_OK)
        
        elif action == 'reject':
            device.reject(request.user, reason)
            return Response({
                'message': 'Device rejected successfully',
                'device': DeviceSerializer(device).data
            }, status=status.HTTP_200_OK)
        
        else:
            return Response(
                {'error': "action must be 'approve' or 'reject'"},
                status=status.HTTP_400_BAD_REQUEST
            )


class PendingDevicesAPIView(ListAPIView):
    """
    Admin endpoint to view all pending device approvals.
    """
    serializer_class = DeviceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    
    def get_queryset(self):
        # Check permission
        if not (self.request.user.is_superuser or self.request.user.has_perm('djangosimplemissionapp.manage_device_approvals')):
            return Device.objects.none()
        
        return Device.objects.filter(is_approved=False).order_by('-created_at')


class RoleListCreateAPIView(ListCreateAPIView):
    queryset = Role.objects.all()
    serializer_class = RoleSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name']

class RoleDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Role.objects.all()
    serializer_class = RoleSerializer
    permission_classes = [IsAuthenticated]

class PermissionListAPIView(views.APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        perms = Permission.objects.all().order_by('codename')
        serializer = PermissionSerializer(perms, many=True)
        return Response([p['codename'] for p in serializer.data])

class RoleCreateAPIView(views.APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        name = request.data.get('name', '').strip().title()
        if not name:
            return Response({'error': 'Role name is required'}, status=status.HTTP_400_BAD_REQUEST)

        perms_codenames = request.data.get('permissions', [])

        role = Role.objects.filter(name__iexact=name).first()
        created = False

        if not role:
            role = Role.objects.create(name=name)
            created = True

        permissions = Permission.objects.filter(codename__in=perms_codenames)
        role.permissions.set(permissions)

        return Response(
            {'message': f'Role {"created" if created else "updated"} with {permissions.count()} permissions'},
            status=status.HTTP_201_CREATED
        )

class AdminChangeUserPasswordView(views.APIView):
    permission_classes = [IsSuperAdmin]

    def put(self, request, pk):
        try:
            user = User.objects.get(pk=pk)
        except User.DoesNotExist:
            raise Http404

        serializer = AdminChangePasswordSerializer(data=request.data)
        if serializer.is_valid():
            user.set_password(serializer.data.get('new_password'))
            user.save()
            return Response({'status': 'password set'}, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ProjectClientListCreateAPIView(ListCreateAPIView):
    queryset = ProjectClient.objects.all()
    serializer_class = ProjectClientSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['company_name', 'contact_person', 'email', 'phone']

class ProjectClientDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectClient.objects.all()
    serializer_class = ProjectClientSerializer
    permission_classes = [IsAuthenticated]

class ProjectBusinessAddressListCreateAPIView(ListCreateAPIView):
    queryset = ProjectBusinessAddress.objects.all()
    serializer_class = ProjectBusinessAddressSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['attention_name', 'city', 'district', 'state', 'pin_code','legal_name', 'projects__name']

    def get_queryset(self):
        queryset = ProjectBusinessAddress.objects.all()
        project_id = self.request.query_params.get('project')
        if project_id:
            queryset = queryset.filter(projects__id=project_id)
        return queryset

class ProjectBusinessAddressDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectBusinessAddress.objects.all()
    serializer_class = ProjectBusinessAddressSerializer
    permission_classes = [IsAuthenticated]

class ClientSummaryListAPIView(ListCreateAPIView):
    queryset = ProjectBusinessAddress.objects.all()
    serializer_class = ClientSummarySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['legal_name', 'city']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
        else:
            serializer = self.get_serializer(queryset, many=True)
            response = Response(serializer.data)

        # Calculate statistics from the filtered queryset
        from django.db.models import Sum
        
        # Base invoice queryset filtered by the clients in the current view
        stats_invoices = Invoice.objects.filter(client_company__in=queryset)
        
        # Apply the same date filters to ensure stats match the displayed criteria
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            stats_invoices = stats_invoices.filter(created_at__date__gte=start_date)
        if end_date:
            stats_invoices = stats_invoices.filter(created_at__date__lte=end_date)

        total_stats = stats_invoices.aggregate(
            total_invoiced=Sum('total_amount'),
            total_paid=Sum('total_paid'),
            total_balance=Sum('balance_due')
        )

        response.data['statistics'] = {
            'total_invoiced': float(total_stats['total_invoiced'] or 0),
            'total_paid': float(total_stats['total_paid'] or 0),
            'total_balance': float(total_stats['total_balance'] or 0)
        }

        return response

    def get_queryset(self):
        from django.db.models import Sum, DecimalField
        from django.db.models.functions import Coalesce

        queryset = ProjectBusinessAddress.objects.all().order_by('-id')

        # Annotate aggregated financial fields for filtering
        queryset = queryset.annotate(
            annotated_balance_due=Coalesce(Sum('invoices__balance_due'), 0, output_field=DecimalField()),
            annotated_total_advance=Coalesce(Sum('advances__amount'), 0, output_field=DecimalField()),
            annotated_remaining_advance=Coalesce(Sum('advances__remaining_amount'), 0, output_field=DecimalField()),
        )

        # Balance status filter
        balance_status = self.request.query_params.get('balance_status', None)
        if balance_status == 'HAS_BALANCE':
            queryset = queryset.filter(annotated_balance_due__gt=0)
        elif balance_status == 'NO_BALANCE':
            queryset = queryset.filter(annotated_balance_due=0)

        # Balance Due range filters
        min_balance = self.request.query_params.get('min_balance')
        max_balance = self.request.query_params.get('max_balance')
        if min_balance:
            queryset = queryset.filter(annotated_balance_due__gte=min_balance)
        if max_balance:
            queryset = queryset.filter(annotated_balance_due__lte=max_balance)

        # Total Advance range filters
        min_advance = self.request.query_params.get('min_advance')
        max_advance = self.request.query_params.get('max_advance')
        if min_advance:
            queryset = queryset.filter(annotated_total_advance__gte=min_advance)
        if max_advance:
            queryset = queryset.filter(annotated_total_advance__lte=max_advance)

        # Remaining Advance range filters
        min_rem_adv = self.request.query_params.get('min_remaining_advance')
        max_rem_adv = self.request.query_params.get('max_remaining_advance')
        if min_rem_adv:
            queryset = queryset.filter(annotated_remaining_advance__gte=min_rem_adv)
        if max_rem_adv:
            queryset = queryset.filter(annotated_remaining_advance__lte=max_rem_adv)

        # Invoice date range filters (filter by invoices created in date range)
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            queryset = queryset.filter(invoices__created_at__date__gte=start_date).distinct()
        if end_date:
            queryset = queryset.filter(invoices__created_at__date__lte=end_date).distinct()

        return queryset

class DomainOrServerThirdPartyServiceProviderListCreateAPIView(ListCreateAPIView):
    queryset = DomainOrServerThirdPartyServiceProvider.objects.all()
    serializer_class = DomainOrServerThirdPartyServiceProviderSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['company_name', 'contact_person', 'email', 'phone']

class DomainOrServerThirdPartyServiceProviderDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = DomainOrServerThirdPartyServiceProvider.objects.all()
    serializer_class = DomainOrServerThirdPartyServiceProviderSerializer
    permission_classes = [IsAuthenticated]

class ProjectDomainListCreateAPIView(ListCreateAPIView):
    queryset = ProjectDomain.objects.all()
    serializer_class = ProjectDomainSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'purchased_from', 'status', 'accrued_by','project__name']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Statistics breakdown
        base_qs = queryset
        stats = {
            'total': base_qs.count(),
            'active': base_qs.filter(status__iexact='Active').count(),
            'pending': base_qs.filter(status__iexact='Pending').count(),
            'expired': base_qs.filter(status__iexact='Expired').count(),
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['statistics'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'results': serializer.data,
            'statistics': stats
        })

    def get_queryset(self):
        from django.utils import timezone
        from django.db.models import Case, When, Value, IntegerField, F
        today = timezone.now().date()
        qs = super().get_queryset()

        # Status filter
        status = self.request.query_params.get('status')
        if status:
            qs = qs.filter(status__iexact=status)

        # Payment Status filter
        payment_status = self.request.query_params.get('payment_status')
        if payment_status:
            qs = qs.filter(payment_status__iexact=payment_status)

        # Invoice Status filter
        invoice_status = self.request.query_params.get('invoice_status')
        if invoice_status:
            qs = qs.filter(invoice_status__iexact=invoice_status)

        # Cost Range filter
        min_cost = self.request.query_params.get('min_cost')
        max_cost = self.request.query_params.get('max_cost')
        if min_cost:
            qs = qs.filter(cost__gte=min_cost)
        if max_cost:
            qs = qs.filter(cost__lte=max_cost)

        # Date Range filter (expiration_date)
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            qs = qs.filter(expiration_date__gte=start_date)
        if end_date:
            qs = qs.filter(expiration_date__lte=end_date)

        return qs.annotate(
            is_expired=Case(
                When(expiration_date__lt=today, then=Value(1)),
                When(status__iexact='Expired', then=Value(1)),
                default=Value(0),
                output_field=IntegerField()
            )
        ).order_by('is_expired', F('expiration_date').asc(nulls_last=True))

class ProjectDomainDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectDomain.objects.all()
    serializer_class = ProjectDomainSerializer
    permission_classes = [IsAuthenticated]    

class ProjectServerListCreateAPIView(ListCreateAPIView):
    queryset = ProjectServer.objects.all()
    serializer_class = ProjectServerSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'server_type', 'purchased_from', 'status', 'accrued_by',"project__name",]

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Statistics breakdown
        base_qs = queryset
        stats = {
            'total': base_qs.count(),
            'active': base_qs.filter(status__iexact='Active').count(),
            'pending': base_qs.filter(status__iexact='Pending').count(),
            'expired': base_qs.filter(status__iexact='Expired').count(),
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['statistics'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'results': serializer.data,
            'statistics': stats
        })

    def get_queryset(self):
        from django.utils import timezone
        from django.db.models import Case, When, Value, IntegerField, F
        today = timezone.now().date()
        qs = super().get_queryset()

        # Status filter
        status = self.request.query_params.get('status')
        if status:
            qs = qs.filter(status__iexact=status)

        # Payment Status filter
        payment_status = self.request.query_params.get('payment_status')
        if payment_status:
            qs = qs.filter(payment_status__iexact=payment_status)

        # Invoice Status filter
        invoice_status = self.request.query_params.get('invoice_status')
        if invoice_status:
            qs = qs.filter(invoice_status__iexact=invoice_status)

        # Cost Range filter
        min_cost = self.request.query_params.get('min_cost')
        max_cost = self.request.query_params.get('max_cost')
        if min_cost:
            qs = qs.filter(cost__gte=min_cost)
        if max_cost:
            qs = qs.filter(cost__lte=max_cost)

        # Date Range filter (expiration_date)
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            qs = qs.filter(expiration_date__gte=start_date)
        if end_date:
            qs = qs.filter(expiration_date__lte=end_date)

        return qs.annotate(
            is_expired=Case(
                When(expiration_date__lt=today, then=Value(1)),
                When(status__iexact='Expired', then=Value(1)),
                default=Value(0),
                output_field=IntegerField()
            )
        ).order_by('is_expired', F('expiration_date').asc(nulls_last=True))

class ProjectServerDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectServer.objects.all()
    serializer_class = ProjectServerSerializer
    permission_classes = [IsAuthenticated]      

class ProjectExbotListCreateAPIView(ListCreateAPIView):
    queryset = ProjectExbot.objects.all()
    serializer_class = ProjectExbotSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['whatsapp_number', 'plan_category', 'status', 'payment_status','project__name']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Statistics breakdown
        base_qs = queryset
        stats = {
            'total': base_qs.count(),
            'active': base_qs.filter(status__iexact='Active').count(),
            'pending': base_qs.filter(status__iexact='Pending').count(),
            'expired': base_qs.filter(status__iexact='Expired').count(),
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['statistics'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'results': serializer.data,
            'statistics': stats
        })

    def get_queryset(self):
        from django.utils import timezone
        from django.db.models import Case, When, Value, IntegerField, F
        today = timezone.now().date()
        qs = super().get_queryset()

        # Status filter
        status = self.request.query_params.get('status')
        if status:
            qs = qs.filter(status__iexact=status)

        # Payment Status filter
        payment_status = self.request.query_params.get('payment_status')
        if payment_status:
            qs = qs.filter(payment_status__iexact=payment_status)

        # Invoice Status filter
        invoice_status = self.request.query_params.get('invoice_status')
        if invoice_status:
            qs = qs.filter(invoice_status__iexact=invoice_status)

        # Rate Range filter
        min_rate = self.request.query_params.get('min_rate')
        max_rate = self.request.query_params.get('max_rate')
        if min_rate:
            qs = qs.filter(plan_rate__gte=min_rate)
        if max_rate:
            qs = qs.filter(plan_rate__lte=max_rate)

        # Plan Category filter
        plan_category = self.request.query_params.get('plan_category')
        if plan_category:
            qs = qs.filter(plan_category__icontains=plan_category)

        # Date Range filter (plan_deactive_date)
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            qs = qs.filter(plan_deactive_date__gte=start_date)
        if end_date:
            qs = qs.filter(plan_deactive_date__lte=end_date)

        return qs.annotate(
            is_expired=Case(
                When(status__iexact='Expired', then=Value(1)),
                default=Value(0),
                output_field=IntegerField()
            )
        ).order_by('is_expired', F('plan_deactive_date').asc(nulls_last=True))

class ProjectExbotDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectExbot.objects.all()
    serializer_class = ProjectExbotSerializer
    permission_classes = [IsAuthenticated]

class ProjectFinanceListCreateAPIView(ListCreateAPIView):
    queryset = ProjectFinance.objects.all()
    serializer_class = ProjectFinanceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['id']

class ProjectFinanceDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectFinance.objects.all()
    serializer_class = ProjectFinanceSerializer
    permission_classes = [IsAuthenticated]          

class TeamListCreateAPIView(ListCreateAPIView):
    serializer_class = TeamSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'team_lead__username']

    def check_permissions(self, request):
        # Check create permission
        if request.method == 'POST' and not request.user.has_perm('djangosimplemissionapp.add_team'):
            self.permission_denied(request, "You don't have permission to create teams.")
        
        # Check view permission
        if request.method == 'GET' and not request.user.has_perm('djangosimplemissionapp.view_team'):
            self.permission_denied(request, "You don't have permission to view teams.")
        
        super().check_permissions(request)

    def get_queryset(self):
        user = self.request.user
        if not user.has_role('SuperAdmin') and not user.has_role('Admin'):
            if not user.has_perm('djangosimplemissionapp.viewall_team ') and user.has_perm('djangosimplemissionapp.viewown_team'):
                return Team.objects.filter(members=user)
        return Team.objects.all()

class TeamDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Team.objects.all()
    serializer_class = TeamSerializer
    permission_classes = [IsAuthenticated]
    
    def check_permissions(self, request):
        # Check delete permission
        if request.method == 'DELETE' and not request.user.has_perm('djangosimplemissionapp.delete_team'):
            self.permission_denied(request, "You don't have permission to delete teams.")
        
        # Check update permission
        if request.method in ['PUT', 'PATCH'] and not request.user.has_perm('djangosimplemissionapp.change_team'):
            self.permission_denied(request, "You don't have permission to edit teams.")
        
        super().check_permissions(request)          

class ProjectTeamListCreateAPIView(ListCreateAPIView):
    queryset = ProjectTeam.objects.all()
    serializer_class = ProjectTeamSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['team__name']

class ProjectTeamDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectTeam.objects.all()
    serializer_class = ProjectTeamSerializer
    permission_classes = [IsAuthenticated]              

class ProjectNatureListCreateAPIView(ListCreateAPIView):
    queryset = ProjectNature.objects.all()
    serializer_class = ProjectNatureSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name']

class ProjectNatureDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectNature.objects.all()
    serializer_class = ProjectNatureSerializer
    permission_classes = [IsAuthenticated]                  

class ProjectListCreateAPIView(ListCreateAPIView):
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'description', 'status', 'project_nature__name', 'project_clients__company_name', 'project_business_addresses__legal_name', 'project_base_informations__name']

    def get_queryset(self):
        user = self.request.user
        queryset = Project.objects.all()

        if user.is_superuser or user.has_role('SuperAdmin'):
            pass
        elif user.has_perm('djangosimplemissionapp.all_projectservicemember') or \
             user.has_perm('djangosimplemissionapp.all_projectteammember'):
            pass
        elif user.has_perm('djangosimplemissionapp.own_projectservicemember') or \
             user.has_perm('djangosimplemissionapp.own_projectteammember'):
            queryset = queryset.filter(
                Q(services__members__employee=user) |
                Q(project_team_members__employee=user)
            ).distinct()
        elif user.has_perm('djangosimplemissionapp.view_project'):
            pass
        else:
            queryset = queryset.none()

        queryset = queryset.order_by('-created_at')
        
        status = self.request.query_params.get('status')
        if status:
            queryset = queryset.filter(status__iexact=status)
            
        nature = self.request.query_params.get('nature')
        if nature:
            queryset = queryset.filter(project_nature__name__icontains=nature)
            
        client = self.request.query_params.get('client')
        if client:
            queryset = queryset.filter(
                Q(project_business_addresses__legal_name__icontains=client) |
                Q(project_clients__company_name__icontains=client)
            ).distinct()
            
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date and end_date:
            queryset = queryset.filter(created_at__date__range=[start_date, end_date])
        elif start_date:
            queryset = queryset.filter(created_at__date__gte=start_date)
        elif end_date:
            queryset = queryset.filter(created_at__date__lte=end_date)
            
        return queryset

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Calculate global stats for the filtered queryset (ignoring pagination)
        stats = {
            'total': queryset.count(),
            'pending': queryset.filter(status__iexact='Pending').count(),
            'completed': queryset.filter(Q(status__iexact='Completed') | Q(status__iexact='Done')).count(),
            'progressing': queryset.exclude(status__iexact='Pending')
                                .exclude(status__iexact='Completed')
                                .exclude(status__iexact='Done')
                                .count()
        }
        
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['stats'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'count': len(serializer.data),
            'results': serializer.data,
            'stats': stats
        })

class ProjectSummaryListAPIView(ListCreateAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSummarySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'status', 'project_base_informations__name']

class ProjectDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if user.is_superuser or user.has_role('SuperAdmin'):
            return queryset

        if user.has_perm('djangosimplemissionapp.all_projectservicemember') or \
           user.has_perm('djangosimplemissionapp.all_projectteammember'):
            return queryset

        if user.has_perm('djangosimplemissionapp.own_projectservicemember') or \
           user.has_perm('djangosimplemissionapp.own_projectteammember'):
            return queryset.filter(
                Q(services__members__employee=user) |
                Q(project_team_members__employee=user)
            ).distinct()

        if user.has_perm('djangosimplemissionapp.view_project') or \
           user.has_perm('djangosimplemissionapp.viewnameonly_projectdomain') or \
           user.has_perm('djangosimplemissionapp.viewnameonly_projectserver'):
            return queryset

        return queryset.none()


class ProjectBaseInformationListCreateAPIView(ListCreateAPIView):
    queryset = ProjectBaseInformation.objects.all()
    serializer_class = ProjectBaseInformationSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'description', 'creator_name']

class ProjectBaseInformationDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectBaseInformation.objects.all()
    serializer_class = ProjectBaseInformationSerializer
    permission_classes = [IsAuthenticated]

class ProjectExcutionListCreateAPIView(ListCreateAPIView):
    queryset = ProjectExcution.objects.all()
    serializer_class = ProjectExcutionSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['id']

class ProjectExcutionDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectExcution.objects.all()
    serializer_class = ProjectExcutionSerializer
    permission_classes = [IsAuthenticated]

class ProjectTeamMemberListCreateAPIView(ListCreateAPIView):
    queryset = ProjectTeamMember.objects.all()
    serializer_class = ProjectTeamMemberSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['employee__username', 'role', 'status']

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if user.is_superuser or user.has_role('SuperAdmin') or user.has_role('Admin'):
            return queryset
            
        if user.has_perm('djangosimplemissionapp.all_projectteammember'):
            return queryset

        if user.has_perm('djangosimplemissionapp.own_projectteammember'):
            return queryset.filter(employee=user)
            
        return queryset.none()

class ProjectTeamMemberDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectTeamMember.objects.all()
    serializer_class = ProjectTeamMemberSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if user.is_superuser or user.has_role('SuperAdmin') or user.has_role('Admin'):
            return queryset
            
        if user.has_perm('djangosimplemissionapp.all_projectteammember'):
            return queryset

        if user.has_perm('djangosimplemissionapp.own_projectteammember'):
            return queryset.filter(employee=user)
            
        return queryset.none()

class ProjectServiceListCreateAPIView(ListCreateAPIView):
    queryset = ProjectService.objects.all()
    serializer_class = ProjectServiceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['description', 'status']

class ProjectServiceDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectService.objects.all()
    serializer_class = ProjectServiceSerializer
    permission_classes = [IsAuthenticated]

class EmployeeDailyActivityListCreateAPIView(ListCreateAPIView):
    # queryset = EmployeeDailyActivity.objects.all()  # Removed in favor of dynamic filtering
    serializer_class = EmployeeDailyActivitySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        can_view_all = self.request.user.has_perm('djangosimplemissionapp.view_all_activities')
        can_view_own = self.request.user.has_perm('djangosimplemissionapp.view_own_activities') or can_view_all
        can_view_team = self.request.user.has_perm('djangosimplemissionapp.viewmeandprojectmember_employeedailyactivity')
        can_view_managers = self.request.user.has_perm('djangosimplemissionapp.viewmanagerallteammember_employeedailyactivity')
        
        if not (can_view_all or can_view_own or can_view_team):
            return EmployeeDailyActivity.objects.none()

        if can_view_all:
            return EmployeeDailyActivity.objects.all()
            
        if can_view_team:
            from django.db.models import Q
            from .models import ProjectTeamMember, ProjectServiceMember
            
            user_projects = ProjectTeamMember.objects.filter(employee=self.request.user).values_list('project_id', flat=True)
            team_members = ProjectTeamMember.objects.filter(project_id__in=user_projects).values_list('employee_id', flat=True)
            
            user_services = ProjectServiceMember.objects.filter(employee=self.request.user).values_list('service_id', flat=True)
            service_members = ProjectServiceMember.objects.filter(service_id__in=user_services).values_list('employee_id', flat=True)

            queryset = EmployeeDailyActivity.objects.filter(
                Q(employee=self.request.user) | 
                Q(employee_id__in=team_members) | 
                Q(employee_id__in=service_members)
            ).distinct()
            
            # Exclude managers if user doesn't have view_manager_all_team_member permission
            if not can_view_managers:
                manager_roles = list(ProjectTeamMember.MANAGER_ROLES) + list(ProjectServiceMember.MANAGER_ROLES)
                manager_ids = ProjectTeamMember.objects.filter(
                    role__in=manager_roles,
                    project_id__in=user_projects
                ).values_list('employee_id', flat=True)
                
                manager_service_ids = ProjectServiceMember.objects.filter(
                    role__in=ProjectServiceMember.MANAGER_ROLES,
                    service_id__in=user_services
                ).values_list('employee_id', flat=True)
                
                manager_ids = set(manager_ids) | set(manager_service_ids)
                queryset = queryset.exclude(employee_id__in=manager_ids)
            
            return queryset

        return EmployeeDailyActivity.objects.filter(employee=self.request.user)
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['description', 'employee__username', 'project__name', 'project_service__name']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Date filtering
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if start_date and end_date:
            queryset = queryset.filter(date__range=[start_date, end_date])
            
        # Employee filtering
        employee_id = request.query_params.get('employee_id')
        if employee_id:
            queryset = queryset.filter(employee_id=employee_id)

        # Activity Report Context
        from django.contrib.auth import get_user_model
        User = get_user_model()
        employee_name = 'All Employees'
        if employee_id:
            employee = User.objects.filter(id=employee_id).first()
            if employee:
                employee_name = employee.get_full_name() or employee.username

        context = {
            'title': 'Employee Activity Report',
            'employee_name': employee_name,
            'date_range': f"{start_date} to {end_date}" if start_date and end_date else "All Time"
        }

        if request.query_params.get('export') == 'pdf':
            from .pdf_utils import generate_activity_pdf
            from django.http import HttpResponse
            buffer = generate_activity_pdf(queryset, context)
            pdf_data = buffer.getvalue()
            filename = f"Employee_Activity_Report_{timezone.now().strftime('%Y%m%d')}.pdf"
            
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Content-Length'] = len(pdf_data)
            response["Access-Control-Allow-Origin"] = "*"
            response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
            response["Access-Control-Allow-Headers"] = "*"
            return response

        elif request.query_params.get('export') == 'docx':
            from .docx_utils import generate_activity_docx
            from django.contrib.auth import get_user_model
            from django.http import HttpResponse
            User = get_user_model()
            
            employee_name = 'All Employees'
            e_id = request.query_params.get('employee_id')
            if e_id:
                employee = User.objects.filter(id=e_id).first()
                if employee:
                    employee_name = employee.username

            context = {
                'title': 'Employee Activity Report',
                'employee_name': employee_name,
                'date_range': f"{start_date} to {end_date}" if start_date and end_date else "All Time"
            }
            buffer = generate_activity_docx(queryset, context)
            docx_data = buffer.getvalue()
            filename = f"Employee_Activity_Report_{timezone.now().strftime('%Y%m%d')}.docx"
            
            response = HttpResponse(docx_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Content-Length'] = len(docx_data)
            response["Access-Control-Allow-Origin"] = "*"
            response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
            response["Access-Control-Allow-Headers"] = "*"
            return response

        # --- SUMMARY LOGIC ---
        # Get the latest activity for each unique assignment (employee + project + service)
        # Since queryset is ordered by -date, -created_at, the first one we see is the latest.
        latest_activities = {}
        for obj in queryset:
            key = (obj.employee_id, obj.project_id, obj.project_service_id)
            if key not in latest_activities:
                latest_activities[key] = obj
                
        temp_serializer = self.get_serializer()
        
        total_allocated_days = 0
        total_remaining_days = 0
        total_overdue_days = 0
        total_used_days = 0
        
        for obj in latest_activities.values():
            total_allocated_days += temp_serializer.get_allocateddays(obj)
            total_remaining_days += temp_serializer.get_remainingdays(obj)
            total_overdue_days += temp_serializer.get_overdue_days(obj)
            total_used_days += temp_serializer.get_totaluseddays(obj)
        
        summary = {
            "total_allocated_days": total_allocated_days,
            "total_remaining_days": total_remaining_days,
            "total_overdue_days": total_overdue_days,
            "total_used_days": total_used_days,
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            
            from collections import OrderedDict
            new_data = OrderedDict()
            new_data['count'] = response.data.get('count')
            new_data['next'] = response.data.get('next')
            new_data['previous'] = response.data.get('previous')
            new_data['summary'] = summary
            new_data['results'] = response.data.get('results')
            response.data = new_data
            
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'summary': summary,
            'results': serializer.data
        })

class EmployeeSpecificActivityListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, employee_id):
        # 0. Permission check
        can_view_all = request.user.has_perm('djangosimplemissionapp.view_all_activities')
        can_view_own = request.user.has_perm('djangosimplemissionapp.view_own_activities') or can_view_all
        can_view_team = request.user.has_perm('djangosimplemissionapp.viewmeandprojectmember_employeedailyactivity')
        
        is_allowed = can_view_all or (can_view_own and request.user.id == int(employee_id))
        
        if not is_allowed and can_view_team:
            from .models import ProjectTeamMember, ProjectServiceMember
            user_projects = ProjectTeamMember.objects.filter(employee=request.user).values_list('project_id', flat=True)
            team_members = ProjectTeamMember.objects.filter(project_id__in=user_projects).values_list('employee_id', flat=True)
            
            user_services = ProjectServiceMember.objects.filter(employee=request.user).values_list('service_id', flat=True)
            service_members = ProjectServiceMember.objects.filter(service_id__in=user_services).values_list('employee_id', flat=True)
            
            if int(employee_id) in team_members or int(employee_id) in service_members:
                is_allowed = True
                
        if not is_allowed:
            return Response({'error': 'Permission denied.'}, status=status.HTTP_403_FORBIDDEN)
        # 1. Fetch activities for the specified employee
        activities = EmployeeDailyActivity.objects.filter(employee_id=employee_id).order_by('-date', '-created_at')
        
        # 2. Support filtering by date range
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if start_date and end_date:
            activities = activities.filter(date__range=[start_date, end_date])

        # 3. Support PDF Export
        if request.query_params.get('export') == 'pdf':
            from .pdf_utils import generate_activity_pdf
            from django.contrib.auth import get_user_model
            User = get_user_model()
            employee = User.objects.filter(id=employee_id).first()
            employee_name = employee.get_full_name() or employee.username if employee else f"User {employee_id}"
            
            context = {
                'title': f'Activity Report for {employee_name}',
                'employee_name': employee_name,
                'date_range': f"{start_date} to {end_date}" if start_date and end_date else "All Time"
            }
            from django.http import HttpResponse
            buffer = generate_activity_pdf(activities, context)
            pdf_data = buffer.getvalue()
            filename = f"Activity_Report_{employee_name}_{timezone.now().strftime('%Y%m%d')}.pdf"
            
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Content-Length'] = len(pdf_data)
            response["Access-Control-Allow-Origin"] = "*"
            response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
            response["Access-Control-Allow-Headers"] = "*"
            return response

        # 4. Standard Paginated Response
        paginator = PageNumberPagination()
        paginator.page_size = int(request.query_params.get('page_size', 10))
        page = paginator.paginate_queryset(activities, request)
        
        serializer = EmployeeDailyActivitySerializer(page, many=True)
        return paginator.get_paginated_response(serializer.data)

class EmployeeWorkDetailsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, employee_id):
        # 0. Permission check for viewing work details
        # For simplicity, we leverage the same 'view_all_activities' / 'view_own_activities' or keep it open for self.
        can_view_all = request.user.has_perm('djangosimplemissionapp.view_all_activities') or \
                       request.user.has_perm('djangosimplemissionapp.view_all_employee_performance')
        can_view_team = request.user.has_perm('djangosimplemissionapp.viewmeandprojectmember_employeedailyactivity')
        
        is_self = request.user.id == int(employee_id)
        is_allowed = can_view_all or is_self
        
        if not is_allowed and can_view_team:
            from .models import ProjectTeamMember, ProjectServiceMember
            user_projects = ProjectTeamMember.objects.filter(employee=request.user).values_list('project_id', flat=True)
            team_members = ProjectTeamMember.objects.filter(project_id__in=user_projects).values_list('employee_id', flat=True)
            
            user_services = ProjectServiceMember.objects.filter(employee=request.user).values_list('service_id', flat=True)
            service_members = ProjectServiceMember.objects.filter(service_id__in=user_services).values_list('employee_id', flat=True)
            
            if int(employee_id) in team_members or int(employee_id) in service_members:
                is_allowed = True

        if not is_allowed:
            return Response({'error': 'Permission denied.'}, status=status.HTTP_403_FORBIDDEN)
        from .models import ProjectTeamMember, ProjectServiceMember, EmployeeDailyActivity
        
        # 1. Projects involved
        project_memberships = ProjectTeamMember.objects.filter(employee_id=employee_id).select_related('project')
        total_projects = project_memberships.values('project').distinct().count()
        
        # 2. Services involved
        service_memberships = ProjectServiceMember.objects.filter(employee_id=employee_id).select_related('service')
        total_services = service_memberships.count()
        active_services = service_memberships.exclude(service__status='Completed').count()
        completed_services = service_memberships.filter(service__status='Completed').count()
        
        # 3. Pending Activities (based on pending_work_percentage > 0)
        pending_activities_count = EmployeeDailyActivity.objects.filter(
            employee_id=employee_id, 
            pending_work_percentage__gt=0
        ).count()

        # 4. List of active projects/services (compact)
        active_items = []
        for pm in project_memberships.exclude(status='Inactive')[:5]:
            days_worked = EmployeeDailyActivity.objects.filter(
                employee_id=employee_id,
                project=pm.project
            ).values('date').distinct().count()
            
            # Fetch project deadline from ProjectExcution
            from .models import ProjectExcution
            deadline = None
            if pm.project:
                exec_info = ProjectExcution.objects.filter(project=pm.project).first()
                if exec_info and exec_info.confirmed_end_date:
                    deadline = exec_info.confirmed_end_date.isoformat()

            active_items.append({
                "type": "Project",
                "name": pm.project.name if pm.project else "Unknown",
                "role": pm.role,
                "days_worked": days_worked,
                "status": pm.status or 'Pending',
                "deadline": deadline
            })
        
        for sm in service_memberships.exclude(status='Inactive')[:5]:
            days_worked = EmployeeDailyActivity.objects.filter(
                employee_id=employee_id,
                project_service=sm.service
            ).values('date').distinct().count()
            
            active_items.append({
                "type": "Service",
                "name": sm.service.name if sm.service else "Unknown",
                "role": sm.role,
                "days_worked": days_worked,
                "status": sm.status or 'Pending',
                "deadline": sm.service.deadline.isoformat() if sm.service and sm.service.deadline else None
            })

        # 5. Recent Activities
        recent_activities = EmployeeDailyActivity.objects.filter(employee_id=employee_id).select_related('project', 'project_service').order_by('-date', '-created_at')[:20]
        activities_data = []
        for act in recent_activities:
            activities_data.append({
                "id": act.id,
                "date": act.date.isoformat() if act.date else None,
                "project_name": act.project.name if act.project else None,
                "service_name": act.project_service.name if act.project_service else None,
                "description": act.description,
                "hours_spent": float(act.hours_spent),
                "target_work_percentage": act.target_work_percentage,
                "pending_work_percentage": act.pending_work_percentage
            })

        data = {
            "total_projects": total_projects,
            "total_services": total_services,
            "active_services": active_services,
            "completed_services": completed_services,
            "pending_activities": pending_activities_count,
            "active_work_list": active_items,
            "recent_activities": activities_data
        }
        
        return Response(data, status=status.HTTP_200_OK)

class EmployeeDailyActivityDetailAPIView(RetrieveUpdateDestroyAPIView):
    # queryset = EmployeeDailyActivity.objects.all()
    serializer_class = EmployeeDailyActivitySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        can_view_all = self.request.user.has_perm('djangosimplemissionapp.view_all_activities')
        can_view_team = self.request.user.has_perm('djangosimplemissionapp.viewmeandprojectmember_employeedailyactivity')

        if can_view_all:
            return EmployeeDailyActivity.objects.all()
            
        if can_view_team:
            from django.db.models import Q
            from .models import ProjectTeamMember, ProjectServiceMember
            
            user_projects = ProjectTeamMember.objects.filter(employee=self.request.user).values_list('project_id', flat=True)
            team_members = ProjectTeamMember.objects.filter(project_id__in=user_projects).values_list('employee_id', flat=True)
            
            user_services = ProjectServiceMember.objects.filter(employee=self.request.user).values_list('service_id', flat=True)
            service_members = ProjectServiceMember.objects.filter(service_id__in=user_services).values_list('employee_id', flat=True)

            return EmployeeDailyActivity.objects.filter(
                Q(employee=self.request.user) | 
                Q(employee_id__in=team_members) | 
                Q(employee_id__in=service_members)
            ).distinct()

        return EmployeeDailyActivity.objects.filter(employee=self.request.user)

class ActivityLogListCreateAPIView(ListCreateAPIView):
    queryset = ActivityLog.objects.all()
    serializer_class = ActivityLogSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['description']

class ActivityLogDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ActivityLog.objects.all()
    serializer_class = ActivityLogSerializer
    permission_classes = [IsAuthenticated]

class ClientInvoiceListAPIView(ListCreateAPIView):
    serializer_class = InvoiceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        client_id = self.kwargs.get('client_id')
        
        # Statistics breakdown for this client
        base_qs = queryset
        stats = {
            'total': base_qs.count(),
            'paid': base_qs.filter(status__iexact='PAID').count(),
            'partial': base_qs.filter(status__iexact='PARTIAL').count(),
            'unpaid': base_qs.filter(status__iexact='UNPAID').count(),
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['statistics'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'results': serializer.data,
            'statistics': stats
        })

    def get_queryset(self):
        client_id = self.kwargs.get('client_id')
        queryset = Invoice.objects.filter(client_company_id=client_id)
        
        # Status filter
        status = self.request.query_params.get('status')
        if status:
            queryset = queryset.filter(status__iexact=status)
            
        # Amount Range filter
        min_amount = self.request.query_params.get('min_amount')
        max_amount = self.request.query_params.get('max_amount')
        if min_amount:
            queryset = queryset.filter(total_amount__gte=min_amount)
        if max_amount:
            queryset = queryset.filter(total_amount__lte=max_amount)
            
        # Date Range filter
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            queryset = queryset.filter(invoice_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(invoice_date__lte=end_date)
            
        return queryset.order_by('-invoice_date')

    def perform_create(self, serializer):
        client_id = self.kwargs.get('client_id')
        serializer.save(client_company_id=client_id)

class InvoiceListCreateAPIView(ListCreateAPIView):
    queryset = Invoice.objects.all()
    serializer_class = InvoiceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    search_fields = ['invoice_number', 'client_company__legal_name']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        
        # Statistics breakdown
        base_qs = queryset
        stats = {
            'total': base_qs.count(),
            'paid': base_qs.filter(status__iexact='PAID').count(),
            'partial': base_qs.filter(status__iexact='PARTIAL').count(),
            'unpaid': base_qs.filter(status__iexact='UNPAID').count(),
        }

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['statistics'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'results': serializer.data,
            'statistics': stats
        })

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Status filter
        status = self.request.query_params.get('status')
        if status:
            queryset = queryset.filter(status__iexact=status)
            
        # Amount Range filter
        min_amount = self.request.query_params.get('min_amount')
        max_amount = self.request.query_params.get('max_amount')
        if min_amount:
            queryset = queryset.filter(total_amount__gte=min_amount)
        if max_amount:
            queryset = queryset.filter(total_amount__lte=max_amount)
            
        # Date Range filter
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            queryset = queryset.filter(invoice_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(invoice_date__lte=end_date)
            
        return queryset.order_by('-invoice_date')

class InvoiceDetailAPIView(RetrieveUpdateDestroyAPIView):
    serializer_class = InvoiceSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        client_id = self.kwargs.get('client_id')
        return Invoice.objects.filter(client_company_id=client_id)

class InvoiceItemListCreateAPIView(ListCreateAPIView):
    queryset = InvoiceItem.objects.all()
    serializer_class = InvoiceItemSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['service_type', 'description']

class InvoiceItemDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = InvoiceItem.objects.all()
    serializer_class = InvoiceItemSerializer
    permission_classes = [IsAuthenticated]

class PaymentListCreateAPIView(ListCreateAPIView):
    queryset = Payment.objects.all()
    serializer_class = PaymentSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['transaction_id', 'payment_method', 'notes']

class PaymentDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Payment.objects.all()
    serializer_class = PaymentSerializer
    permission_classes = [IsAuthenticated]

class InvoicePaymentListCreateAPIView(ListCreateAPIView):
    """
    Nested endpoint: /api/project-business-addresses/<client_id>/invoices/<invoice_id>/payments/
    Handle GET (list payments for this invoice) and POST (add new payment).
    """
    serializer_class = PaymentSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination

    def get_queryset(self):
        invoice_id = self.kwargs.get('pk')
        client_id = self.kwargs.get('client_id')
        return Payment.objects.filter(invoice_id=invoice_id, invoice__client_company_id=client_id)

    def perform_create(self, serializer):
        invoice_id = self.kwargs.get('pk')
        client_id = self.kwargs.get('client_id')
        try:
            invoice = Invoice.objects.get(id=invoice_id, client_company_id=client_id)
        except Invoice.DoesNotExist:
            raise Http404("Invoice not found for this client")
        
        serializer.save(invoice=invoice)

class InvoicePaymentDetailAPIView(RetrieveUpdateDestroyAPIView):
    """
    Nested endpoint: /api/project-business-addresses/<client_id>/invoices/<invoice_id>/payments/<payment_id>/
    Handle specific payment management.
    """
    serializer_class = PaymentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        invoice_id = self.kwargs.get('invoice_pk')
        client_id = self.kwargs.get('client_id')
        return Payment.objects.filter(invoice_id=invoice_id, invoice__client_company_id=client_id)

class ApplyAdvanceCreditView(APIView):
    """
    Manually apply a specific amount of advance credit to an invoice.
    Endpoint: /api/project-business-addresses/<client_id>/invoices/<invoice_id>/apply-advance/
    """
    permission_classes = [IsAuthenticated]

    def post(self, request, client_id, pk):
        try:
            invoice = Invoice.objects.get(id=pk, client_company_id=client_id)
        except Invoice.DoesNotExist:
            return Response({"error": "Invoice not found for this client"}, status=status.HTTP_404_NOT_FOUND)

        amount = request.data.get('amount')
        if not amount:
            return Response({"error": "Amount is required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            amount = Decimal(str(amount))
        except:
            return Response({"error": "Invalid amount format"}, status=status.HTTP_400_BAD_REQUEST)

        if amount <= 0:
            return Response({"error": "Amount must be greater than zero"}, status=status.HTTP_400_BAD_REQUEST)

        if amount > invoice.balance_due:
            return Response({"error": f"Amount exceeds invoice balance ({invoice.balance_due})"}, status=status.HTTP_400_BAD_REQUEST)

        from .services import apply_advances_to_invoice
        
        # Check if they have enough advance total
        from .models import ClientAdvance
        from django.db.models import Sum
        available = ClientAdvance.objects.filter(client=invoice.client_company, remaining_amount__gt=0).aggregate(Sum('remaining_amount'))['remaining_amount__sum'] or 0
        
        if amount > available:
            return Response({"error": f"Insufficient advance credit (Available: {available})"}, status=status.HTTP_400_BAD_REQUEST)

        apply_advances_to_invoice(invoice, limit_amount=amount)
        
        return Response({
            "message": f"Successfully applied {amount} from advance credit.",
            "new_balance": invoice.balance_due
        }, status=status.HTTP_200_OK)

class ActivityExceedCommentListCreateAPIView(ListCreateAPIView):
    serializer_class = ActivityExceedCommentSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['comment', 'commented_by__username']

    def get_queryset(self):
        queryset = ActivityExceedComment.objects.all()
        activity_id = self.request.query_params.get('activity')
        if activity_id:
            queryset = queryset.filter(activity_id=activity_id)
        return queryset.order_by('-created_at')

class ActivityExceedCommentDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ActivityExceedComment.objects.all()
    serializer_class = ActivityExceedCommentSerializer
    permission_classes = [IsAuthenticated]

class NotificationListCreateAPIView(ListCreateAPIView):
    serializer_class = NotificationSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['message', 'user__username']

    def get_queryset(self):
        user = self.request.user
        qs = Notification.objects.filter(user=user)
        
        # If user is superuser, show all their notifications
        if user.is_superuser:
            return qs.order_by('-created_at')
            
        # Check for specific granular notification permissions
        can_view_all = user.has_perm('djangosimplemissionapp.view_notification')
        can_view_server = user.has_perm('djangosimplemissionapp.view_server_notifications')
        can_view_domain = user.has_perm('djangosimplemissionapp.view_domain_notifications')
        can_view_exbot = user.has_perm('djangosimplemissionapp.view_exbot_notifications')

        # If they don't have the global "view_notification" permission, 
        # we filter by the specific types they ARE allowed to see.
        if not can_view_all:
            allowed_types = []
            if can_view_server: allowed_types.append('server_alert')
            if can_view_domain: allowed_types.append('domain_alert')
            if can_view_exbot: allowed_types.append('exbot_alert')
            
            # If they have specific permissions, filter by those types
            if allowed_types:
                qs = qs.filter(notification_type__in=allowed_types)
            else:
                # If they have NONE of the above, they only see notifications 
                # that don't have these specific alert types (e.g. system messages)
                qs = qs.exclude(notification_type__in=['server_alert', 'domain_alert', 'exbot_alert'])
                
        return qs.order_by('-created_at')

class NotificationDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Notification.objects.all()
    serializer_class = NotificationSerializer
    permission_classes = [IsAuthenticated]

class UnreadNotificationCountAPIView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        count = Notification.objects.filter(user=request.user, is_read=False).count()
        return Response({'unread_count': count})

class MarkAllNotificationsReadAPIView(views.APIView):
    permission_classes = [IsAuthenticated]
    def put(self, request):
        Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
        return Response({'message': 'All notifications marked as read'}, status=status.HTTP_200_OK)

class EmployeeLeaveListCreateAPIView(ListCreateAPIView):
    serializer_class = EmployeeLeaveSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['employee__username', 'status', 'description']

    def get_queryset(self):
        user = self.request.user

        if user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_employeeleave'):
            queryset = EmployeeLeave.objects.all().order_by('-start_date')
            # allow admins to filter by employee
            employee_id = self.request.query_params.get('employee')
            if employee_id:
                queryset = queryset.filter(employee_id=employee_id)
            
            project_id = self.request.query_params.get('project')
            if project_id:
                from .models import ProjectTeamMember
                # Find all employees that belong to this project
                team_employees = ProjectTeamMember.objects.filter(
                    project_id=project_id, 
                    status__in=['Pending', 'Progressing']
                ).values_list('employee_id', flat=True)
                queryset = queryset.filter(employee_id__in=team_employees)
        else:
            queryset = EmployeeLeave.objects.filter(employee=user).order_by('-start_date')

        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        status_filter = self.request.query_params.get('status')

        if start_date:
            queryset = queryset.filter(start_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(end_date__lte=end_date)
            
        if status_filter:
            queryset = queryset.filter(status__iexact=status_filter)
            
        return queryset

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        # Compute stats on the full filtered queryset (before pagination)
        from django.db.models import Count, Q
        stats = queryset.aggregate(
            total=Count('id'),
            approved=Count('id', filter=Q(status__iexact='approved')),
            rejected=Count('id', filter=Q(status__iexact='rejected')),
            cancelled=Count('id', filter=Q(status__iexact='cancelled')),
            pending=Count('id', filter=Q(status__iexact='pending') | Q(status__isnull=True) | Q(status='')),
        )

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['stats'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({'results': serializer.data, 'stats': stats})

class EmployeeLeaveDetailAPIView(RetrieveUpdateDestroyAPIView):
    serializer_class = EmployeeLeaveSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user

        if user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_employeeleave'):
            return EmployeeLeave.objects.all()

        return EmployeeLeave.objects.filter(employee=user)

    def perform_update(self, serializer):
        user = self.request.user
        new_status = serializer.validated_data.get('status')

        if new_status:
            normalized_status = str(new_status).strip().lower()

            if normalized_status in ['approved', 'approve']:
                if not user.is_superuser and not user.has_perm('djangosimplemissionapp.approve_employeeleave'):
                    raise PermissionDenied("You do not have permission to approve leave.")
                serializer.save(approved_by=user)
                return

            if normalized_status in ['rejected', 'reject']:
                if not user.is_superuser and not user.has_perm('djangosimplemissionapp.reject_employeeleave'):
                    raise PermissionDenied("You do not have permission to reject leave.")
                serializer.save(approved_by=user)
                return

        serializer.save()

class CompanyListCreateAPIView(ListCreateAPIView):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['id']

class CompanyDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer
    permission_classes = [IsAuthenticated]

class CompanyProfileListCreateAPIView(ListCreateAPIView):
    queryset = CompanyProfile.objects.all()
    serializer_class = CompanyProfileSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['company_name', 'company_type', 'email', 'phone', 'address']

class CompanyProfileDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = CompanyProfile.objects.all()
    serializer_class = CompanyProfileSerializer
    permission_classes = [IsAuthenticated]

class EmployeeSalarySummaryListAPIView(ListAPIView):
    queryset = User.objects.filter(salaries__isnull=False).distinct()
    serializer_class = EmployeeSalarySummarySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['username', 'first_name', 'last_name']

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # User role based filtering
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_salary')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_salary') or can_view_all
        
        if not (can_view_all or can_view_own):
            # If user has neither permission, deny access
            return queryset.none()
        
        if not can_view_all:
            # If can only view own, filter to their records
            queryset = queryset.filter(id=user.id)
        
        # Filter by specific user/employee
        user_filter = self.request.query_params.get('user_id')
        if user_filter:
            queryset = queryset.filter(id=user_filter)
        
        # Filter by salary status
        status_filter = self.request.query_params.get('status')
        if status_filter:
            queryset = queryset.filter(salaries__status__iexact=status_filter).distinct()
        
        # Filter by minimum paid count
        from django.db.models import Count, Q
        min_paid_count = self.request.query_params.get('min_paid_count')
        if min_paid_count:
            try:
                min_paid = int(min_paid_count)
                queryset = queryset.annotate(
                    paid_count=Count('salaries', filter=Q(salaries__status='Paid'))
                ).filter(paid_count__gte=min_paid)
            except (ValueError, TypeError):
                pass
        
        # Filter by minimum unpaid count
        min_unpaid_count = self.request.query_params.get('min_unpaid_count')
        if min_unpaid_count:
            try:
                min_unpaid = int(min_unpaid_count)
                queryset = queryset.annotate(
                    unpaid_count=Count('salaries', filter=Q(salaries__status='Unpaid'))
                ).filter(unpaid_count__gte=min_unpaid)
            except (ValueError, TypeError):
                pass
        
        return queryset

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
        else:
            serializer = self.get_serializer(queryset, many=True)
            response = Response(serializer.data)

        # Calculate statistics based on filtered employees' salaries
        from django.db.models import Sum, Count, Case, When, IntegerField, Q
        
        # Get filtered salary records
        filtered_salaries = Salary.objects.filter(employee__in=queryset)
        
        # Apply status filter if provided
        status_filter = self.request.query_params.get('status')
        if status_filter:
            filtered_salaries = filtered_salaries.filter(status__iexact=status_filter)
        
        # Apply date range filters if provided
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            filtered_salaries = filtered_salaries.filter(created_at__date__gte=start_date)
        if end_date:
            filtered_salaries = filtered_salaries.filter(created_at__date__lte=end_date)
        
        stats = filtered_salaries.aggregate(
            total_basic=Sum('basic'),
            paid_count=Count(Case(When(status='Paid', then=1), output_field=IntegerField())),
            unpaid_count=Count(Case(When(status='Unpaid', then=1), output_field=IntegerField()))
        )

        response.data['statistics'] = {
            'total_basic': float(stats['total_basic'] or 0),
            'paid_count': stats['paid_count'] or 0,
            'unpaid_count': stats['unpaid_count'] or 0
        }

        return response

class SalaryListCreateAPIView(ListCreateAPIView):
    queryset = Salary.objects.all()
    serializer_class = SalarySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['employee__user__username', 'status']

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
        else:
            serializer = self.get_serializer(queryset, many=True)
            response = Response(serializer.data)

        # Calculate statistics from the filtered queryset
        from django.db.models import Sum, Count, Case, When, IntegerField
        stats = queryset.aggregate(
            total_basic=Sum('basic'),
            paid_count=Count(Case(When(status='Paid', then=1), output_field=IntegerField())),
            unpaid_count=Count(Case(When(status='Unpaid', then=1), output_field=IntegerField()))
        )

        response.data['statistics'] = {
            'total_basic': float(stats['total_basic'] or 0),
            'paid_count': stats['paid_count'] or 0,
            'unpaid_count': stats['unpaid_count'] or 0
        }

        return response

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # User role based filtering
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_salary')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_salary') or can_view_all
        
        if not (can_view_all or can_view_own):
            # If user has neither permission, deny access
            return queryset.none()
        
        if not can_view_all:
            # If can only view own, filter to their records
            queryset = queryset.filter(employee=user)
        
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        employee_id = self.request.query_params.get('employee')
        
        if start_date:
            queryset = queryset.filter(start_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(end_date__lte=end_date)
        if can_view_all and employee_id:
            queryset = queryset.filter(employee_id=employee_id)
            
        return queryset

class SalaryDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Salary.objects.all()
    serializer_class = SalarySerializer
    permission_classes = [IsAuthenticated]

class AttendanceListCreateAPIView(ListCreateAPIView):
    queryset = Attendance.objects.all().order_by('-date')
    serializer_class = AttendanceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['employee__username', 'status', 'approval_status']

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # User role based filtering
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_attendance')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_attendance') or can_view_all
        
        if not (can_view_all or can_view_own):
            # If user has neither permission, deny access
            return queryset.none()
        
        if not can_view_all:
            # If can only view own, filter to their records
            queryset = queryset.filter(employee=user)
        else:
            # allow admins to filter by employee
            employee_id = self.request.query_params.get('employee')
            if employee_id:
                queryset = queryset.filter(employee_id=employee_id)
            
            project_id = self.request.query_params.get('project')
            if project_id:
                from .models import ProjectTeamMember
                # Find all employees that belong to this project
                team_employees = ProjectTeamMember.objects.filter(
                    project_id=project_id, 
                    status__in=['Pending', 'Progressing']
                ).values_list('employee_id', flat=True)
                queryset = queryset.filter(employee_id__in=team_employees)

        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        date = self.request.query_params.get('date')
        status_filter = self.request.query_params.get('status')
        approval_status = self.request.query_params.get('approval_status')
        check_in_from = self.request.query_params.get('check_in_from')
        check_in_to = self.request.query_params.get('check_in_to')

        if date:
            queryset = queryset.filter(date=date)
        else:
            if start_date:
                queryset = queryset.filter(date__gte=start_date)
            if end_date:
                queryset = queryset.filter(date__lte=end_date)
                
        if status_filter:
            queryset = queryset.filter(status=status_filter)
            
        if approval_status:
            queryset = queryset.filter(approval_status=approval_status)
        
        # Filter by check_in time range
        if check_in_from:
            # Ensure check_in_from is in HH:MM:SS format
            try:
                if check_in_from and len(check_in_from.split(':')) == 2:  # HH:MM format
                    check_in_from = f"{check_in_from}:00"
                print(f"DEBUG: Filtering by check_in_from = {check_in_from}")
                queryset = queryset.filter(check_in__gte=check_in_from)
            except Exception as e:
                print(f"Error parsing check_in_from: {e}")
                
        if check_in_to:
            # Ensure check_in_to is in HH:MM:SS format
            try:
                if check_in_to and len(check_in_to.split(':')) == 2:  # HH:MM format
                    check_in_to = f"{check_in_to}:00"
                print(f"DEBUG: Filtering by check_in_to = {check_in_to}")
                queryset = queryset.filter(check_in__lte=check_in_to)
            except Exception as e:
                print(f"Error parsing check_in_to: {e}")
        
        print(f"DEBUG: Final queryset count = {queryset.count()}")
        return queryset

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        # Compute stats on the full filtered queryset (before pagination)
        from django.db.models import Count, Q
        stats = queryset.aggregate(
            total=Count('id'),
            present=Count('id', filter=Q(status='Present')),
            absent=Count('id', filter=Q(status='Absent')),
            half_day=Count('id', filter=Q(status='Half Day')),
            late=Count('id', filter=Q(status='Late')),
            approved=Count('id', filter=Q(approval_status='Approved')),
            rejected=Count('id', filter=Q(approval_status='Rejected')),
            pending=Count('id', filter=Q(approval_status='Pending') | Q(approval_status__isnull=True)),
        )

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
            response.data['stats'] = stats
            return response

        serializer = self.get_serializer(queryset, many=True)
        return Response({'results': serializer.data, 'stats': stats})

class AttendanceDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Attendance.objects.all()
    serializer_class = AttendanceSerializer
    permission_classes = [IsAuthenticated]

class AttendanceExportAPIView(APIView):
    """Export attendance records to Excel or Word format"""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        Export filtered attendance records
        Query params: export_format (excel/word), status, approval_status, start_date, end_date, employee, check_in_from, check_in_to
        """
        from datetime import datetime
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from django.http import FileResponse
        import io

        # Get format from request (renamed to export_format to avoid DRF content negotiation)
        export_format = request.query_params.get('export_format', 'excel').lower()
        if export_format not in ['excel', 'word']:
            return Response({'error': 'Invalid export_format. Use "excel" or "word"'}, status=400)

        # Get filtered queryset using same logic as list view
        user = request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_attendance')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_attendance') or can_view_all

        if not (can_view_all or can_view_own):
            return Response({'error': 'You do not have permission to export attendance'}, status=403)

        queryset = Attendance.objects.all().order_by('-date')

        if not can_view_all:
            queryset = queryset.filter(employee=user)
        else:
            employee_id = request.query_params.get('employee')
            if employee_id:
                queryset = queryset.filter(employee_id=employee_id)

        # Apply filters
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        status_filter = request.query_params.get('status')
        approval_status = request.query_params.get('approval_status')
        check_in_from = request.query_params.get('check_in_from')
        check_in_to = request.query_params.get('check_in_to')

        if start_date:
            queryset = queryset.filter(date__gte=start_date)
        if end_date:
            queryset = queryset.filter(date__lte=end_date)
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if approval_status:
            queryset = queryset.filter(approval_status=approval_status)
        if check_in_from:
            if len(check_in_from.split(':')) == 2:
                check_in_from = f"{check_in_from}:00"
            queryset = queryset.filter(check_in__gte=check_in_from)
        if check_in_to:
            if len(check_in_to.split(':')) == 2:
                check_in_to = f"{check_in_to}:00"
            queryset = queryset.filter(check_in__lte=check_in_to)

        if export_format == 'excel':
            return self._export_to_excel(queryset)
        else:
            return self._export_to_word(queryset)

    def _export_to_excel(self, queryset):
        """Export to Excel format"""
        from datetime import datetime
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from django.http import FileResponse
        import io

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Attendance"

        # Define styles
        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        center_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Headers
        headers = ['#', 'Employee Name', 'Employee ID', 'Date', 'Check In', 'Check Out', 'Working Hours', 'Status', 'Approval Status']
        worksheet.append(headers)

        # Format header row
        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Data
        for idx, record in enumerate(queryset, 1):
            employee_name = f"{record.employee.first_name} {record.employee.last_name}".strip() or record.employee.username
            check_in = str(record.check_in).split('.')[0] if record.check_in else '—'
            check_out = str(record.check_out).split('.')[0] if record.check_out else '—'
            working_hours = record.get_total_working_hours() if hasattr(record, 'get_total_working_hours') else 0

            row = [
                idx,
                employee_name,
                record.employee_id,
                record.date.strftime('%Y-%m-%d'),
                check_in,
                check_out,
                f"{working_hours}h" if working_hours else "—",
                record.status,
                record.approval_status or 'Pending'
            ]
            worksheet.append(row)

            # Format data cells
            for cell in worksheet[idx + 1]:
                cell.border = thin_border
                cell.alignment = center_alignment if cell.column in [1, 4, 6, 7, 8, 9] else Alignment(horizontal="left", vertical="center")

        # Adjust column widths
        worksheet.column_dimensions['A'].width = 5
        worksheet.column_dimensions['B'].width = 20
        worksheet.column_dimensions['C'].width = 15
        worksheet.column_dimensions['D'].width = 12
        worksheet.column_dimensions['E'].width = 12
        worksheet.column_dimensions['F'].width = 12
        worksheet.column_dimensions['G'].width = 15
        worksheet.column_dimensions['H'].width = 15
        worksheet.column_dimensions['I'].width = 15

        # Save to bytes
        excel_file = io.BytesIO()
        workbook.save(excel_file)
        excel_file.seek(0)

        filename = f"Attendance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return FileResponse(
            excel_file,
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    def _export_to_word(self, queryset):
        """Export to Word format"""
        from datetime import datetime
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from django.http import FileResponse
        import io

        document = Document()

        # Add title
        title = document.add_paragraph()
        title_run = title.add_run("Attendance Report")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(54, 96, 146)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()  # Blank line

        # Create table
        table = document.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'

        # Add header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Employee Name', 'Employee ID', 'Date', 'Check In', 'Check Out', 'Working Hours', 'Status', 'Approval Status']
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            # Style header
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '366092')
            header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)

        # Add data rows
        for idx, record in enumerate(queryset, 1):
            row = table.add_row()
            cells = row.cells

            employee_name = f"{record.employee.first_name} {record.employee.last_name}".strip() or record.employee.username
            check_in = str(record.check_in).split('.')[0] if record.check_in else '—'
            check_out = str(record.check_out).split('.')[0] if record.check_out else '—'
            working_hours = record.get_total_working_hours() if hasattr(record, 'get_total_working_hours') else 0

            data = [
                str(idx),
                employee_name,
                str(record.employee_id),
                record.date.strftime('%Y-%m-%d'),
                check_in,
                check_out,
                f"{working_hours}h" if working_hours else "—",
                record.status,
                record.approval_status or 'Pending'
            ]

            for cell_idx, cell_text in enumerate(data):
                cells[cell_idx].text = cell_text
                for paragraph in cells[cell_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        word_file = io.BytesIO()
        document.save(word_file)
        word_file.seek(0)

        filename = f"Attendance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return FileResponse(
            word_file,
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

class EmployeeLeaveExportAPIView(APIView):
    """Export employee leaves to Excel or Word format"""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        Export filtered leave records
        Query params: export_format (excel/word), employee, start_date, end_date, status
        """
        from datetime import datetime
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from django.http import FileResponse
        import io

        # Get format from request
        export_format = request.query_params.get('export_format', 'excel').lower()
        if export_format not in ['excel', 'word']:
            return Response({'error': 'Invalid export_format. Use "excel" or "word"'}, status=400)

        # Get filtered queryset using same logic as list view
        user = request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_employeeleave')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_employeeleave') or can_view_all

        # For leaves, even basic permission allows viewing own leaves
        if not (can_view_all or can_view_own or user.has_perm('djangosimplemissionapp.viewall_employeeleave')):
            return Response({'error': 'You do not have permission to export leaves'}, status=403)

        queryset = EmployeeLeave.objects.all().order_by('-start_date')

        if not can_view_all:
            queryset = queryset.filter(employee=user)
        else:
            employee_id = request.query_params.get('employee')
            if employee_id:
                queryset = queryset.filter(employee_id=employee_id)

        # Apply filters
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        status_filter = request.query_params.get('status')
        project_id = request.query_params.get('project')

        if start_date:
            queryset = queryset.filter(start_date__gte=start_date)
        if end_date:
            queryset = queryset.filter(end_date__lte=end_date)
        if status_filter:
            queryset = queryset.filter(status__iexact=status_filter)
        if project_id:
            from .models import ProjectTeamMember
            team_employees = ProjectTeamMember.objects.filter(
                project_id=project_id,
                status__in=['Pending', 'Progressing']
            ).values_list('employee_id', flat=True)
            queryset = queryset.filter(employee_id__in=team_employees)

        if export_format == 'excel':
            return self._export_to_excel(queryset)
        else:
            return self._export_to_word(queryset)

    def _export_to_excel(self, queryset):
        """Export to Excel format"""
        from datetime import datetime
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from django.http import FileResponse
        import io

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Leaves"

        # Define styles
        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        center_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Headers
        headers = ['#', 'Employee Name', 'Employee ID', 'Start Date', 'End Date', 'Leave Type', 'Duration (Days)', 'Status', 'Description', 'Approved By']
        worksheet.append(headers)

        # Format header row
        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Data
        for idx, record in enumerate(queryset, 1):
            employee_name = f"{record.employee.first_name} {record.employee.last_name}".strip() or record.employee.username
            approved_by_name = f"{record.approved_by.first_name} {record.approved_by.last_name}".strip() if record.approved_by else '—'
            
            # Calculate duration
            duration = (record.end_date - record.start_date).days + 1
            if record.leave_type == 'Half Day':
                duration = 0.5

            row = [
                idx,
                employee_name,
                record.employee_id,
                record.start_date.strftime('%Y-%m-%d'),
                record.end_date.strftime('%Y-%m-%d'),
                record.leave_type or 'Full Day',
                f"{duration}" if duration > 0 else "—",
                record.status or 'Pending',
                record.description or '—',
                approved_by_name
            ]
            worksheet.append(row)

            # Format data cells
            for cell in worksheet[idx + 1]:
                cell.border = thin_border
                cell.alignment = center_alignment if cell.column in [1, 4, 5, 6, 7, 8, 10] else Alignment(horizontal="left", vertical="center")

        # Adjust column widths
        worksheet.column_dimensions['A'].width = 5
        worksheet.column_dimensions['B'].width = 20
        worksheet.column_dimensions['C'].width = 15
        worksheet.column_dimensions['D'].width = 12
        worksheet.column_dimensions['E'].width = 12
        worksheet.column_dimensions['F'].width = 15
        worksheet.column_dimensions['G'].width = 15
        worksheet.column_dimensions['H'].width = 12
        worksheet.column_dimensions['I'].width = 25
        worksheet.column_dimensions['J'].width = 15

        # Save to bytes
        excel_file = io.BytesIO()
        workbook.save(excel_file)
        excel_file.seek(0)

        filename = f"Leaves_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return FileResponse(
            excel_file,
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    def _export_to_word(self, queryset):
        """Export to Word format"""
        from datetime import datetime
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from django.http import FileResponse
        import io

        document = Document()

        # Add title
        title = document.add_paragraph()
        title_run = title.add_run("Employee Leaves Report")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(54, 96, 146)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()  # Blank line

        # Create table
        table = document.add_table(rows=1, cols=10)
        table.style = 'Light Grid Accent 1'

        # Add header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Employee Name', 'Employee ID', 'Start Date', 'End Date', 'Leave Type', 'Duration (Days)', 'Status', 'Description', 'Approved By']
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            # Style header
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '366092')
            header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)

        # Add data rows
        for idx, record in enumerate(queryset, 1):
            row = table.add_row()
            cells = row.cells

            employee_name = f"{record.employee.first_name} {record.employee.last_name}".strip() or record.employee.username
            approved_by_name = f"{record.approved_by.first_name} {record.approved_by.last_name}".strip() if record.approved_by else '—'
            
            # Calculate duration
            duration = (record.end_date - record.start_date).days + 1
            if record.leave_type == 'Half Day':
                duration = 0.5

            data = [
                str(idx),
                employee_name,
                str(record.employee_id),
                record.start_date.strftime('%Y-%m-%d'),
                record.end_date.strftime('%Y-%m-%d'),
                record.leave_type or 'Full Day',
                f"{duration}" if duration > 0 else "—",
                record.status or 'Pending',
                record.description or '—',
                approved_by_name
            ]

            for cell_idx, cell_text in enumerate(data):
                cells[cell_idx].text = cell_text
                for paragraph in cells[cell_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        word_file = io.BytesIO()
        document.save(word_file)
        word_file.seek(0)

        filename = f"Leaves_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return FileResponse(
            word_file,
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

class EmployeeListCreateAPIView(ListCreateAPIView):
    queryset = Employee.objects.all()
    serializer_class = EmployeeSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['user__username', 'employee_id', 'department']

class EmployeeDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Employee.objects.all()
    serializer_class = EmployeeSerializer
    permission_classes = [IsAuthenticated]

class UserSalaryListCreateAPIView(ListCreateAPIView):
    serializer_class = UserSalarySerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['user__username', 'user__first_name', 'user__last_name']

    def get_queryset(self):
        qs = UserSalary.objects.all().order_by('-effective_date', '-created_at')
        
        # User role based filtering
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_usersalary')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_usersalary') or can_view_all
        
        if not (can_view_all or can_view_own):
            # If user has neither permission, deny access
            return qs.none()
        
        if not can_view_all:
            # If can only view own, filter to their records
            qs = qs.filter(user=user)
        
        user_id = self.request.query_params.get('user')
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        min_salary = self.request.query_params.get('min_salary')
        max_salary = self.request.query_params.get('max_salary')
        department = self.request.query_params.get('department')
        
        if can_view_all and user_id:
            qs = qs.filter(user_id=user_id)
        if start_date:
            qs = qs.filter(effective_date__gte=start_date)
        if end_date:
            qs = qs.filter(effective_date__lte=end_date)
        if min_salary:
            qs = qs.filter(base_salary__gte=float(min_salary))
        if max_salary:
            qs = qs.filter(base_salary__lte=float(max_salary))
        if department:
            qs = qs.filter(user__employee__department=department)
            
        return qs

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            response = self.get_paginated_response(serializer.data)
        else:
            serializer = self.get_serializer(queryset, many=True)
            response = Response(serializer.data)

        # Calculate statistics from the filtered queryset
        from django.db.models import Sum, Count, Avg
        stats = queryset.aggregate(
            total_count=Count('id'),
            total_base_salary=Sum('base_salary'),
            average_salary=Avg('base_salary'),
            unique_employees=Count('user', distinct=True)
        )

        response.data = {
            'results': response.data if isinstance(response.data, list) else response.data.get('results', []),
            'statistics': {
                'total_count': stats['total_count'] or 0,
                'total_base_salary': float(stats['total_base_salary'] or 0),
                'average_salary': float(stats['average_salary'] or 0),
                'unique_employees': stats['unique_employees'] or 0
            }
        }
        
        if hasattr(response, 'get') and callable(getattr(response, 'get')):
            # Preserve pagination info if it exists
            if 'count' in response.data:
                response.data['count'] = queryset.count()

        return response

class UserSalaryDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = UserSalary.objects.all()
    serializer_class = UserSalarySerializer
    permission_classes = [IsAuthenticated]

    def retrieve(self, request, *args, **kwargs):
        """Override retrieve to include next salary day in response"""
        response = super().retrieve(request, *args, **kwargs)
        
        user_salary = self.get_object()
        next_salary_day = self.calculate_next_salary_day(user_salary)
        
        response.data['next_salary_day'] = next_salary_day
        
        return response

    def calculate_next_salary_day(self, user_salary):
        """
        Calculate the next salary day based on joining_date and working_days cycle.
        Returns the end date of the current cycle if joining_date exists.
        """
        from datetime import timedelta
        from django.utils import timezone
        
        if not user_salary.joining_date:
            return None
        
        # Import the helper functions from models
        from .models import get_cycle_end_date, get_cycle_for_date
        
        today = timezone.now().date()
        working_days = user_salary.working_days or 26
        
        # Get the current cycle for today
        start_date, end_date = get_cycle_for_date(user_salary.joining_date, today, working_days)
        
        # If today is within a cycle, return the end date of that cycle
        if end_date:
            return end_date.isoformat()
        
        # Fallback: calculate next cycle end date
        return get_cycle_end_date(today, working_days).isoformat()


# ============================================================================
# SALARY INCREMENT API VIEWS
# ============================================================================

class SalaryIncrementListCreateAPIView(ListCreateAPIView):
    """
    List and create salary increments.
    POST creates a new salary increment and automatically updates UserSalary.
    """
    serializer_class = SalaryIncrementSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['employee__username', 'remarks']

    def get_queryset(self):
        qs = SalaryIncrement.objects.all().order_by('-effective_date', '-created_at')
        
        # User role based filtering
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_salaryincrement')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_salaryincrement') or can_view_all
        
        if not (can_view_all or can_view_own):
            # If user has neither permission, deny access
            return qs.none()
        
        if not can_view_all:
            # If can only view own, filter to their records
            qs = qs.filter(employee=user)
        
        employee_id = self.request.query_params.get('employee')
        if can_view_all and employee_id:
            qs = qs.filter(employee_id=employee_id)
        return qs

    def perform_create(self, serializer):
        """
        Set created_by to current user when creating increment.
        Check permissions: can_view_all can create for any employee, 
        can_view_own can only create for themselves.
        """
        user = self.request.user
        can_view_all = user.is_superuser or user.has_perm('djangosimplemissionapp.viewall_salaryincrement')
        can_view_own = user.has_perm('djangosimplemissionapp.viewown_salaryincrement') or can_view_all
        
        if not (can_view_all or can_view_own):
            raise PermissionDenied("You don't have permission to create salary increments.")
        
        # Get the employee from the request data
        employee_id = serializer.validated_data.get('employee').id if 'employee' in serializer.validated_data else self.request.data.get('employee')
        
        # If user only has view_own permission, they can only create for themselves
        if can_view_own and not can_view_all:
            if int(employee_id) != user.id:
                raise PermissionDenied("You can only create salary increments for yourself.")
        
        serializer.save(created_by=user)


class SalaryIncrementDetailAPIView(RetrieveUpdateDestroyAPIView):
    """
    Retrieve, update, or delete salary increment records.
    """
    queryset = SalaryIncrement.objects.all()
    serializer_class = SalaryIncrementSerializer
    permission_classes = [IsAuthenticated]


class OtherIncomeListCreateAPIView(ListCreateAPIView):

    queryset = OtherIncome.objects.all()
    serializer_class = OtherIncomeSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['title', 'notes']

class OtherIncomeDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = OtherIncome.objects.all()
    serializer_class = OtherIncomeSerializer
    permission_classes = [IsAuthenticated]

class UserDesignationsAPIView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        designations = User.objects.exclude(designation__isnull=True).exclude(designation='').values_list('designation', flat=True).distinct()
        return Response(sorted(list(designations)))

class OtherExpenseListCreateAPIView(ListCreateAPIView):
    queryset = OtherExpense.objects.all()
    serializer_class = OtherExpenseSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['title', 'notes']

class OtherExpenseDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = OtherExpense.objects.all()
    serializer_class = OtherExpenseSerializer
    permission_classes = [IsAuthenticated]

class ProjectDocumentListCreateAPIView(ListCreateAPIView):
    queryset = ProjectDocument.objects.all()
    serializer_class = ProjectDocumentSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'description']

class ProjectDocumentDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = ProjectDocument.objects.all()
    serializer_class = ProjectDocumentSerializer
    permission_classes = [IsAuthenticated]

class InvoicePDFView(APIView):  
    permission_classes = [IsAuthenticated]

    def get(self, request, client_id, pk):
        try:
            invoice = Invoice.objects.get(pk=pk, client_company_id=client_id)
        except Invoice.DoesNotExist:
            return Response({"error": "Invoice not found for this client"}, status=status.HTTP_404_NOT_FOUND)

        from django.http import HttpResponse
        buffer = generate_invoice_pdf(invoice)
        pdf_data = buffer.getvalue()
        filename = f"Invoice_{invoice.invoice_number or invoice.id}.pdf"
        
        response = HttpResponse(pdf_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = len(pdf_data)
        
        # Manual Security/CORS headers for cPanel stability
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        response["Access-Control-Allow-Headers"] = "*"
        
        return response

class ClientAdvanceListAPIView(ListCreateAPIView):
    serializer_class = ClientAdvanceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination

    def get_queryset(self):
        client_id = self.kwargs.get('client_id')
        return ClientAdvance.objects.filter(client_id=client_id).order_by('created_at')

    def perform_create(self, serializer):
        client_id = self.kwargs.get('client_id')
        serializer.save(client_id=client_id)

class ClientAdvanceDetailAPIView(RetrieveUpdateDestroyAPIView):
    serializer_class = ClientAdvanceSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        client_id = self.kwargs.get('client_id')
        return ClientAdvance.objects.filter(client_id=client_id)

class EmployeePerformanceAPIView(APIView):
    """
    Returns performance analytics for the logged-in employee (or a specific employee
    if employee_id is passed as a query param by an admin).
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        can_view_all = request.user.has_perm('djangosimplemissionapp.view_all_employee_performance')
        can_view_own = request.user.has_perm('djangosimplemissionapp.view_own_employee_performance') or can_view_all

        if not (can_view_all or can_view_own):
            return Response({'error': 'You do not have permission to view employee performance.'}, status=status.HTTP_403_FORBIDDEN)

        employee_id = request.query_params.get('employee_id')
        if employee_id:
            if not can_view_all:
                return Response({'error': 'Permission denied to view other employee performance.'}, status=status.HTTP_403_FORBIDDEN)
            try:
                employee = User.objects.get(pk=employee_id)
            except User.DoesNotExist:
                raise Http404
        else:
            employee = request.user

        from .models import ProjectTeamMember, ProjectServiceMember
        
        project_team_memberships = ProjectTeamMember.objects.filter(employee=employee).select_related('project')
        service_team_memberships = ProjectServiceMember.objects.filter(employee=employee).select_related('service', 'service__project')

        from django.utils import timezone
        current_date = timezone.now().date()

        def get_target_score(start, current, allocated):
            if not start or not allocated or allocated <= 0:
                return 0
            days_passed = (current - start).days
            if days_passed < 0:
                return 0
            return round((days_passed / allocated) * 100, 2)

        project_team_data = []
        for m in project_team_memberships:
            # Calculate elapsed days from start_date to current_date
            days_used = (current_date - m.start_date).days if m.start_date else 0
            days_remaining = max(0, m.allocated_days - days_used)
            usage_percentage = round((days_used / m.allocated_days * 100), 2) if m.allocated_days > 0 else 0
            project_team_data.append({
                'project_name': m.project.name if m.project else 'Unknown Project',
                'status': m.status,
                'start_date': m.start_date,
                'allocation': m.allocated_days,
                'days_used': days_used,
                'days_remaining': days_remaining,
                'usage_percentage': usage_percentage,
                'current_date': current_date,
                'end_date': m.end_date,
                'note': m.notes,
                'is_over_allocated': days_used > m.allocated_days,
                'is_complete': m.status == 'Completed',
            })
            
        service_team_data = []
        for m in service_team_memberships:
            p_name = m.service.project.name if m.service and m.service.project else 'Independent Service'
            s_name = m.service.name if m.service else 'Unknown Service'
            # Calculate elapsed days from start_date to current_date
            days_used = (current_date - m.start_date).days if m.start_date else 0
            days_remaining = max(0, m.allocated_days - days_used)
            usage_percentage = round((days_used / m.allocated_days * 100), 2) if m.allocated_days > 0 else 0
            service_team_data.append({
                'service_name': f"{p_name} -> {s_name}",
                'status': m.status,
                'start_date': m.start_date,
                'allocation': m.allocated_days,
                'days_used': days_used,
                'days_remaining': days_remaining,
                'usage_percentage': usage_percentage,
                'current_date': current_date,
                'end_date': m.end_date,
                'note': m.notes,
                'is_over_allocated': days_used > m.allocated_days,
                'is_complete': m.status == 'Completed',
            })

        pending_count = 0
        progressing_count = 0
        completed_count = 0

        for item in project_team_data:
            item_status = item.get('status')
            if item_status == 'Pending':
                pending_count += 1
            elif item_status == 'Progressing':
                progressing_count += 1
            elif item_status == 'Completed':
                completed_count += 1

        for item in service_team_data:
            item_status = item.get('status')
            if item_status == 'Pending':
                pending_count += 1
            elif item_status == 'Progressing':
                progressing_count += 1
            elif item_status == 'Completed':
                completed_count += 1

        data = {
            'employee_id': employee.id,
            'employee_name': employee.get_full_name() or employee.username,
            'pending_total': pending_count,
            'completed_total': completed_count,
            'progressing_total': progressing_count,
            'total_committed_project_count': len(project_team_data),
            'total_committed_project_team': project_team_data,
            'total_committed_service_count': len(service_team_data),
            'total_committed_service_team': service_team_data,
        }

        return Response(data, status=status.HTTP_200_OK)

    def _get_stats(self, employee):
        from django.db.models import Sum

        memberships = ProjectTeamMember.objects.filter(employee=employee).select_related('project')
        service_memberships = ProjectServiceMember.objects.filter(employee=employee).select_related('service', 'service__project')

        total_allocated = memberships.aggregate(s=Sum('allocated_days'))['s'] or 0
        total_actual = memberships.aggregate(s=Sum('actual_days_spent'))['s'] or 0
        svc_allocated = service_memberships.aggregate(s=Sum('allocated_days'))['s'] or 0
        svc_actual = service_memberships.aggregate(s=Sum('actual_days'))['s'] or 0

        total_allocated += svc_allocated
        total_actual += svc_actual

        time_saved = max(0, total_allocated - total_actual)
        time_overrun = max(0, total_actual - total_allocated)

        perf_pct = round((total_actual / total_allocated) * 100, 1) if total_allocated else 0

        # Combine completed counts
        completed_count = memberships.filter(status='Completed').count() + service_memberships.filter(status='Completed').count()
        progress_count = memberships.filter(status__in=['Completed', 'Progressing']).count() + service_memberships.filter(status__in=['Completed', 'Progressing']).count()
        
        reliability = round((completed_count / max(progress_count, 1)) * 100, 1)

        # Productivity score (0–10): lower actual/allocated is better
        productivity = round(max(0, 10 - (time_overrun / max(total_allocated, 1)) * 10), 2)

        if reliability >= 90 and perf_pct <= 110:
            grade = 'A'
        elif reliability >= 75:
            grade = 'B'
        elif reliability >= 60:
            grade = 'C'
        else:
            grade = 'D'

        if productivity >= 7:
            risk = 'Low'
        elif productivity >= 4:
            risk = 'Medium'
        else:
            risk = 'High'

        projects_list = []
        # Add Project Memberships
        for m in memberships[:15]:
            projects_list.append({
                'id': m.project.id if m.project else None,
                'name': m.project.name if m.project else 'Unknown Project',
                'role': m.role,
                'status': m.status,
                'allocated_days': m.allocated_days,
                'actual_days_spent': m.actual_days_spent,
                'type': 'Project'
            })
            
        # Add Service Memberships
        for m in service_memberships[:15]:
            p_name = m.service.project.name if m.service and m.service.project else 'Independent Service'
            s_name = m.service.name if m.service else 'Unknown Service'
            projects_list.append({
                'id': m.service.id if m.service else None,
                'name': f"{p_name} ➔ {s_name}",
                'role': m.role,
                'status': m.status,
                'allocated_days': m.allocated_days,
                'actual_days_spent': m.actual_days,
                'type': 'Service'
            })

        return {
            'employee_id': employee.id,
            'employee_username': employee.username,
            'projects_count': progress_count,  # Total combined active/completed tasks
            'services_count': service_memberships.count(),
            'total_allocated_days': total_allocated,
            'total_actual_days': total_actual,
            'time_saved_days': time_saved,
            'time_overrun_days': time_overrun,
            'overall_performance_percentage': perf_pct,
            'delivery_reliability_percent': reliability,
            'productivity_score': productivity,
            'performance_grade': grade,
            'performance_risk': risk,
            'overall_status': 'On Track' if risk == 'Low' else ('At Risk' if risk == 'Medium' else 'Delayed'),
            'projects': projects_list,
        }

class TeamPerformanceAPIView(APIView):
    """
    Returns aggregated performance analytics for a team.
    Team leads see their own team; admins can specify team_id.
    """
    permission_classes = [IsAuthenticated]

    def _get_team_stats(self, team):
        from .models import ProjectTeam, ProjectServiceTeam, ProjectTeamMember, ProjectServiceMember
        from django.utils import timezone
        current_date = timezone.now().date()
        
        members = team.members.all()
        member_count = members.count()
        
        # 1. Projects assigned specifically to this Team
        team_project_allocations = ProjectTeam.objects.filter(team=team).select_related('project')
        projects_data = []
        project_ids = set()
        
        p_pending = 0
        p_progressing = 0
        p_completed = 0
        
        for allocation in team_project_allocations:
            if allocation.project and allocation.project.id not in project_ids:
                overused = False
                over_days = 0
                remain_days = 0
                check_date = allocation.deadline or allocation.end_date
                
                if check_date:
                    if allocation.status == 'Completed':
                        if allocation.actual_end_date and allocation.actual_end_date > check_date:
                            overused = True
                            over_days = (allocation.actual_end_date - check_date).days
                    else:
                        days_diff = (check_date - current_date).days
                        if days_diff < 0:
                            overused = True
                            over_days = abs(days_diff)
                        else:
                            remain_days = days_diff
                
                projects_data.append({
                    'name': allocation.project.name,
                    'status': allocation.status, # Use the status from ProjectTeam
                    'start_date': allocation.start_date,
                    'end_date': allocation.end_date,
                    'deadline': allocation.deadline,
                    'actual_end_date': allocation.actual_end_date,
                    'current_date': current_date,
                    'overused': overused,
                    'over_days': over_days,
                    'remain_days': remain_days
                })
                project_ids.add(allocation.project.id)
            
            # Status counts from members linked to this team allocation
            for m in allocation.members.all():
                if m.status == 'Pending': p_pending += 1
                elif m.status == 'Progressing': p_progressing += 1
                elif m.status == 'Completed': p_completed += 1

        # 2. Services assigned specifically to this Team
        team_service_allocations = ProjectServiceTeam.objects.filter(team=team).select_related('service', 'service__project')
        services_data = []
        service_ids = set()
        
        s_pending = 0
        s_progressing = 0
        s_completed = 0
        
        for allocation in team_service_allocations:
            svc = allocation.service
            if svc:
                if svc.id not in service_ids:
                    p_name = svc.project.name if svc.project else 'Independent'
                    overused = False
                    over_days = 0
                    remain_days = 0
                    # Use deadline for overuse check if available, otherwise end_date
                    check_date = allocation.deadline or allocation.end_date
                    
                    if check_date:
                        if allocation.status == 'Completed':
                            if allocation.actual_end_date and allocation.actual_end_date > check_date:
                                overused = True
                                over_days = (allocation.actual_end_date - check_date).days
                        else:
                            days_diff = (check_date - current_date).days
                            if days_diff < 0:
                                overused = True
                                over_days = abs(days_diff)
                            else:
                                remain_days = days_diff
                        
                    services_data.append({
                        'name': f"{p_name} ➔ {svc.name}",
                        'status': allocation.status, # Use the status from ProjectServiceTeam
                        'start_date': allocation.start_date,
                        'end_date': allocation.end_date,
                        'deadline': allocation.deadline,
                        'actual_end_date': allocation.actual_end_date,
                        'current_date': current_date,
                        'overused': overused,
                        'over_days': over_days,
                        'remain_days': remain_days
                    })
                    service_ids.add(svc.id)
                
                # For services, we check ProjectServiceMember matching this service and this team's members
                s_mems = ProjectServiceMember.objects.filter(service=svc, employee__in=members)
                s_pending += s_mems.filter(status='Pending').count()
                s_progressing += s_mems.filter(status='Progressing').count()
                s_completed += s_mems.filter(status='Completed').count()

        return {
            'team_id': team.id,
            'team_name': team.name,
            'member_count': member_count,
            'member_names': [m.username for m in members],
            'team_projects_count': len(projects_data),
            'projects': projects_data,
            'team_service_count': len(services_data),
            'services': services_data,
            # Internal fields for aggregation
            '_pending': p_pending + s_pending,
            '_completed': p_completed + s_completed,
            '_progressing': p_progressing + s_progressing,
        }

    def get(self, request):
        # 1. Permission checks
        role_names_upper = [r.upper() for r in request.user.role_names]
        is_admin_flag = any(r in ['SUPERADMIN', 'ADMIN'] for r in role_names_upper)
        
        def check_perm(codename):
            if request.user.is_superuser or is_admin_flag: return True
            if request.user.has_perm(f'djangosimplemissionapp.{codename}'): return True
            if request.user.role and request.user.role.permissions.filter(codename=codename).exists():
                return True
            return False

        can_view_all = check_perm('view_all_team_performance')
        can_view_own = check_perm('view_own_team_performance') or can_view_all

        if not (can_view_all or can_view_own):
            return Response({'error': 'You do not have permission to view team performance.'}, status=status.HTTP_403_FORBIDDEN)

        # 2. Determine base queryset
        if can_view_all:
            queryset = Team.objects.all()
        else:
            # Users with "Own Team" only see teams where they are lead or member
            queryset = Team.objects.filter(
                Q(team_lead=request.user) | Q(members=request.user)
            ).distinct()

        # 3. Handle specific team drill-down if requested
        team_id = request.query_params.get('team_id')
        if team_id:
            queryset = queryset.filter(pk=team_id)
            if not queryset.exists():
                return Response({'error': 'Team not found or access denied.'}, status=status.HTTP_404_NOT_FOUND)

        # 4. Slice queryset for pagination BEFORE building stats (performance)
        try:
            page = int(request.query_params.get('page', 1))
            page_size = int(request.query_params.get('page_size', 10))
        except (ValueError, TypeError):
            page = 1
            page_size = 10

        page_size = max(1, min(page_size, 100))  # clamp between 1–100

        total_count = queryset.count()
        total_pages = (total_count + page_size - 1) // page_size

        start = (page - 1) * page_size
        end = start + page_size
        paged_queryset = queryset[start:end]

        # 5. Process stats only for this page
        all_stats = [self._get_team_stats(t) for t in paged_queryset]

        total_pending = sum(s.pop('_pending', 0) for s in all_stats)
        total_completed = sum(s.pop('_completed', 0) for s in all_stats)
        total_inprogress = sum(s.pop('_progressing', 0) for s in all_stats)

        # 6. Handle empty cases for better UX
        if not all_stats and not can_view_all:
            return Response({'error': 'No team found for your account.'}, status=status.HTTP_404_NOT_FOUND)

        # Build next/previous URLs
        base_url = request.build_absolute_uri(request.path)

        def build_url(p):
            return f"{base_url}?page={p}&page_size={page_size}"

        return Response({
            'count': total_count,
            'total_pages': total_pages,
            'current_page': page,
            'page_size': page_size,
            'next': build_url(page + 1) if page < total_pages else None,
            'previous': build_url(page - 1) if page > 1 else None,
            'total_teams': total_count,
            'total_pending': total_pending,
            'total_completed': total_completed,
            'total_inprogress': total_inprogress,
            'teams': all_stats,
        }, status=status.HTTP_200_OK)

class LeadViewSet(viewsets.ModelViewSet):
    queryset = Lead.objects.all().order_by('-created_at')
    serializer_class = LeadSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['company_name', 'contact_person', 'contact_number', 'location', 'lead_source']
    ordering_fields = ['created_at', 'next_followup_date', 'interest_level', 'conversion_status']

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        
        can_view_all = user.has_perm('djangosimplemissionapp.view_all_leads')
        can_view_own = user.has_perm('djangosimplemissionapp.view_own_leads') or can_view_all
        
        if not can_view_all:
            if can_view_own:
                qs = qs.filter(assigned_to=user)
            else:
                # If they have neither, they see nothing
                qs = qs.none()
        
        today = timezone.now().date()
        
        # Subqueries for specific follow-up dates
        pending_followups = FollowUp.objects.filter(lead=OuterRef('pk'), status='pending')
        
        # Annotate queryset with nearest upcoming and latest overdue dates
        qs = qs.annotate(
            nearest_upcoming_date=Subquery(pending_followups.filter(followup_date__gte=today).order_by('followup_date').values('followup_date')[:1]),
            latest_overdue_date=Subquery(pending_followups.filter(followup_date__lt=today).order_by('-followup_date').values('followup_date')[:1]),
            current_status=Subquery(FollowUp.objects.filter(lead=OuterRef('pk')).order_by('-created_at').values('conversion_status')[:1]),
            current_interest=Subquery(FollowUp.objects.filter(lead=OuterRef('pk')).order_by('-created_at').values('interest_level')[:1])
        )
        
        # Additional filtering
        interest = self.request.query_params.get('interest_level')
        if interest:
            if interest == 'warm':
                qs = qs.filter(Q(current_interest=interest) | Q(current_interest__isnull=True))
            else:
                qs = qs.filter(current_interest=interest)
        
        lead_conversion_status = self.request.query_params.get('conversion_status')
        if lead_conversion_status:
            if lead_conversion_status == 'new':
                qs = qs.filter(Q(current_status=lead_conversion_status) | Q(current_status__isnull=True))
            else:
                qs = qs.filter(current_status=lead_conversion_status)

        assigned_to = self.request.query_params.get('assigned_to')
        if assigned_to:
            qs = qs.filter(assigned_to_id=assigned_to)

        # Date-based filtering uses the most relevant date (upcoming if exists, else overdue)
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            qs = qs.filter(Q(nearest_upcoming_date__gte=start_date) | Q(latest_overdue_date__gte=start_date))
        if end_date:
            qs = qs.filter(Q(nearest_upcoming_date__lte=end_date) | Q(latest_overdue_date__lte=end_date))
        
        # Upcoming filter — leads with at least one pending upcoming follow-up
        upcoming = self.request.query_params.get('upcoming')
        if upcoming in ('true', '1', 'yes'):
            qs = qs.filter(nearest_upcoming_date__isnull=False)

        # Overdue filter — leads with no upcoming follow-ups but at least one overdue follow-up
        overdue = self.request.query_params.get('overdue')
        if overdue in ('true', '1', 'yes'):
            qs = qs.filter(nearest_upcoming_date__isnull=True, latest_overdue_date__isnull=False).exclude(current_status__in=['closed', 'denied'])

        return qs

class FollowUpViewSet(viewsets.ModelViewSet):
    queryset = FollowUp.objects.all().order_by('-created_at')

    serializer_class = FollowUpSerializer
    permission_classes = [IsAuthenticated]

    def perform_create(self, serializer):
        # When a new interaction is logged, mark all previous pending follow-ups as completed
        lead = serializer.validated_data.get('lead')
        if lead:
            FollowUp.objects.filter(lead=lead, status='pending').update(status='completed')
        serializer.save()

    def get_queryset(self):
        qs = super().get_queryset()
        
        # Filter by lead
        lead_id = self.request.query_params.get('lead')
        if lead_id:
            qs = qs.filter(lead_id=lead_id)
            
        # Filter by specific date
        date = self.request.query_params.get('date')
        if date:
            qs = qs.filter(followup_date=date)

        # Upcoming filter — follow-ups where followup_date >= today
        upcoming = self.request.query_params.get('upcoming')
        if upcoming in ('true', '1', 'yes'):
            from django.utils import timezone
            today = timezone.now().date()
            qs = qs.filter(followup_date__gte=today)

        # Filter by assignment (users only see followups for leads assigned to them)
        user = self.request.user
        if not user.is_superuser:
            qs = qs.filter(lead__assigned_to=user)

        return qs.select_related('lead')

class LeadDashboardStatsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        can_view_all = user.has_perm('djangosimplemissionapp.view_all_leads')
        can_view_own = user.has_perm('djangosimplemissionapp.view_own_leads') or can_view_all

        base_qs = Lead.objects.all()
        if not can_view_all:
            if can_view_own:
                base_qs = base_qs.filter(assigned_to=user)
            else:
                base_qs = base_qs.none()

        from django.db.models import Count, Subquery, OuterRef
        from django.utils import timezone
        today = timezone.now().date()
        pending_followups = FollowUp.objects.filter(lead=OuterRef('pk'), status='pending')
        
        # Annotate queryset with nearest upcoming and latest overdue dates
        base_qs = base_qs.annotate(
            nearest_upcoming_date=Subquery(pending_followups.filter(followup_date__gte=today).order_by('followup_date').values('followup_date')[:1]),
            latest_overdue_date=Subquery(pending_followups.filter(followup_date__lt=today).order_by('-followup_date').values('followup_date')[:1]),
            current_status=Subquery(FollowUp.objects.filter(lead=OuterRef('pk')).order_by('-created_at').values('conversion_status')[:1]),
            current_interest=Subquery(FollowUp.objects.filter(lead=OuterRef('pk')).order_by('-created_at').values('interest_level')[:1])
        )

        # Count by Interest Level
        interest_stats = base_qs.values('current_interest').annotate(count=Count('id'))
        
        # Count by Conversion Status
        status_stats = base_qs.values('current_status').annotate(count=Count('id'))
        
        # Follow-up stats using the new logic
        upcoming_followups = base_qs.filter(nearest_upcoming_date__isnull=False).count()
        overdue_followups = base_qs.filter(nearest_upcoming_date__isnull=True, latest_overdue_date__isnull=False).exclude(current_status__in=['closed', 'denied']).count()

        # Format output, defaulting empty statuses to 'warm' and 'new' to match LeadSerializer logic
        formatted_interest = {}
        for item in interest_stats:
            key = item['current_interest'] or 'warm'
            formatted_interest[key] = formatted_interest.get(key, 0) + item['count']
            
        formatted_status = {}
        for item in status_stats:
            key = item['current_status'] or 'new'
            formatted_status[key] = formatted_status.get(key, 0) + item['count']

        return Response({
            'interest_stats': formatted_interest,
            'status_stats': formatted_status,
            'upcoming_followups': upcoming_followups,
            'overdue_followups': overdue_followups,
            'total_leads': base_qs.count()
        })

class ScheduleListCreateAPIView(ListCreateAPIView):
    serializer_class = ScheduleSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['title', 'description', 'assigned_to__username']

    def get_queryset(self):
        user = self.request.user
        queryset = Schedule.objects.all()
        
        # Check if user has management roles (SuperAdmin, Admin, HR)
        user_roles = user.role_names
        is_management = any(role in ['SuperAdmin', 'Admin', 'HR'] for role in user_roles)

        # Filter by assignment, creator, or role unless management or superuser
        if not user.is_superuser and not is_management:
            from django.db.models import Q
            # User sees schedules assigned to them, created by them, or assigned to their current role
            queryset = queryset.filter(
                Q(assigned_to=user) | 
                Q(created_by=user) | 
                Q(assigned_role=user.role)
            )

            
        # Filter by completion status (default to hide completed)
        show_completed = self.request.query_params.get('show_completed', 'false').lower() == 'true'
        if not show_completed:
            queryset = queryset.filter(is_completed=False)
            
        # Filter by date
        date_param = self.request.query_params.get('date', None)
        if date_param:
            queryset = queryset.filter(schedule_date=date_param)

        # Filter by user
        user_param = self.request.query_params.get('assigned_to', None)
        if user_param:
            queryset = queryset.filter(assigned_to_id=user_param)

        # Filter by role
        role_param = self.request.query_params.get('assigned_role', None)
        if role_param:
            queryset = queryset.filter(assigned_role_id=role_param)

            
        return queryset.order_by('-created_at')

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

class ScheduleDetailAPIView(RetrieveUpdateDestroyAPIView):
    queryset = Schedule.objects.all()
    serializer_class = ScheduleSerializer
    permission_classes = [IsAuthenticated]


class LoginHistoryListAPIView(ListCreateAPIView):
    """
    API endpoint to view user login history.
    - GET: View all logins (admins) or own logins (regular users)
    - Filters available: user_id, ip_address, device_type, login_status, date_range
    """
    serializer_class = LoginUserDetailsSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['ip_address', 'device_name', 'browser_name', 'os_name', 'country', 'city']
    ordering_fields = ['login_time', 'logout_time', 'is_suspicious']
    ordering = ['-login_time']
    
    def get_queryset(self):
        user = self.request.user
        queryset = LoginUserDetails.objects.all()
        
        # Regular users can only see their own logins
        if not user.is_superuser and not user.has_perm('djangosimplemissionapp.view_all_logins'):
            queryset = queryset.filter(user=user)
        
        # Filter by user_id if provided (only admins can see other users)
        user_id = self.request.query_params.get('user_id')
        if user_id:
            if user.is_superuser or user.has_perm('djangosimplemissionapp.view_all_logins'):
                queryset = queryset.filter(user_id=user_id)
            elif int(user_id) != user.id:
                # Non-admin users cannot view other users' logins
                return LoginUserDetails.objects.none()
        
        # Filter by device_type
        device_type = self.request.query_params.get('device_type')
        if device_type:
            queryset = queryset.filter(device_type=device_type)
        
        # Filter by login_status
        login_status = self.request.query_params.get('login_status')
        if login_status:
            queryset = queryset.filter(login_status=login_status)
        
        # Filter by IP address
        ip_address = self.request.query_params.get('ip_address')
        if ip_address:
            queryset = queryset.filter(ip_address=ip_address)
        
        # Filter by suspicious flag
        is_suspicious = self.request.query_params.get('is_suspicious')
        if is_suspicious:
            queryset = queryset.filter(is_suspicious=is_suspicious.lower() == 'true')
        
        # Filter by date range
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            from datetime import datetime
            try:
                start_dt = datetime.fromisoformat(start_date)
                queryset = queryset.filter(login_time__gte=start_dt)
            except:
                pass
        if end_date:
            from datetime import datetime
            try:
                end_dt = datetime.fromisoformat(end_date)
                queryset = queryset.filter(login_time__lte=end_dt)
            except:
                pass
        
        return queryset


class LoginHistoryDetailAPIView(RetrieveUpdateDestroyAPIView):
    """
    API endpoint to view, update, or delete a specific login record.
    Users can only access their own records unless they are admins.
    """
    queryset = LoginUserDetails.objects.all()
    serializer_class = LoginUserDetailsSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        user = self.request.user
        queryset = LoginUserDetails.objects.all()
        
        # Regular users can only see their own logins
        if not user.is_superuser and not user.has_perm('djangosimplemissionapp.view_all_logins'):
            queryset = queryset.filter(user=user)
        
        return queryset
    
    def perform_update(self, serializer):
        """
        Allow updating logout_time and suspicious flag.
        """
        instance = self.get_object()
        
        # Ensure users can only update their own records or admins
        if instance.user != self.request.user and not self.request.user.is_superuser:
            raise PermissionDenied("You cannot update other users' login records.")
        
        # Calculate session duration if logout_time is being set
        if serializer.validated_data.get('logout_time') and not instance.logout_time:
            from datetime import datetime
            logout_time = serializer.validated_data['logout_time']
            if logout_time:
                instance.session_duration = logout_time - instance.login_time
        
        serializer.save()


class CurrentUserLoginHistoryAPIView(views.APIView):
    """
    API endpoint for current user to view their own login history.
    Provides a simple endpoint without filtering requirements.
    """
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """Get current user's login history with pagination"""
        user = request.user
        page = int(request.query_params.get('page', 1))
        per_page = int(request.query_params.get('per_page', 10))
        
        logins = LoginUserDetails.objects.filter(user=user).order_by('-login_time')
        
        # Manual pagination
        start = (page - 1) * per_page
        end = start + per_page
        
        serializer = LoginUserDetailsSerializer(logins[start:end], many=True)
        
        return Response({
            'count': logins.count(),
            'page': page,
            'per_page': per_page,
            'total_pages': (logins.count() + per_page - 1) // per_page,
            'results': serializer.data
        })


class SuspiciousLoginsAPIView(views.APIView):
    """
    API endpoint to view suspicious login attempts.
    Only accessible to admins and superusers.
    """
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """Get all suspicious logins"""
        user = request.user
        
        # Check permission
        if not user.is_superuser and not user.has_perm('djangosimplemissionapp.view_suspicious_logins'):
            raise PermissionDenied("You don't have permission to view suspicious logins.")
        
        suspicious_logins = LoginUserDetails.objects.filter(is_suspicious=True).order_by('-login_time')
        
        page = int(request.query_params.get('page', 1))
        per_page = int(request.query_params.get('per_page', 10))
        
        start = (page - 1) * per_page
        end = start + per_page
        
        serializer = LoginUserDetailsSerializer(suspicious_logins[start:end], many=True)
        
        return Response({
            'count': suspicious_logins.count(),
            'page': page,
            'per_page': per_page,
            'total_pages': (suspicious_logins.count() + per_page - 1) // per_page,
            'results': serializer.data
        })


class LoginUserDetailsListCreateAPIView(ListCreateAPIView):
    """
    List all login details or create a new login record.
    """
    queryset = LoginUserDetails.objects.all()
    serializer_class = LoginUserDetailsSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter]
    search_fields = ['user__username', 'ip_address', 'device_name', 'browser_name']

    def delete(self, request, *args, **kwargs):
        """Delete all login user details. Restricted to superuser or users with delete_loginuserdetails permission."""
        user = request.user
        if not (user.is_superuser or user.has_perm('djangosimplemissionapp.delete_loginuserdetails')):
            return Response({'error': 'You do not have permission to delete all login records.'}, status=status.HTTP_403_FORBIDDEN)
        
        # Delete all login records
        deleted_count, _ = LoginUserDetails.objects.all().delete()
        return Response({'message': f'Successfully deleted all {deleted_count} login records.'}, status=status.HTTP_200_OK)

    def get_queryset(self):
        queryset = LoginUserDetails.objects.all().order_by('-login_time')
        
        # Filter by user if specified
        user_id = self.request.query_params.get('user_id')
        if user_id:
            queryset = queryset.filter(user_id=user_id)
        
        # Filter by login status (case-insensitive)
        login_status = self.request.query_params.get('login_status')
        if login_status:
            queryset = queryset.filter(login_status__iexact=login_status)
        
        # Filter by device type (case-insensitive)
        device_type = self.request.query_params.get('device_type')
        if device_type:
            queryset = queryset.filter(device_type__iexact=device_type)
        
        # Filter by suspicious flag
        is_suspicious = self.request.query_params.get('is_suspicious')
        if is_suspicious:
            queryset = queryset.filter(is_suspicious=is_suspicious.lower() == 'true')
        
        # Filter by year if specified
        year = self.request.query_params.get('year')
        if year:
            queryset = queryset.filter(login_time__year=year)
            
        # Filter by start date if specified
        start_date = self.request.query_params.get('start_date')
        if start_date:
            queryset = queryset.filter(login_time__date__gte=start_date)
            
        # Filter by end date if specified
        end_date = self.request.query_params.get('end_date')
        if end_date:
            queryset = queryset.filter(login_time__date__lte=end_date)
        
        return queryset


class LoginUserDetailsDetailAPIView(RetrieveUpdateDestroyAPIView):
    """
    Retrieve, update, or delete a specific login detail record.
    """
    queryset = LoginUserDetails.objects.all()
    serializer_class = LoginUserDetailsSerializer
    permission_classes = [IsAuthenticated]


class LoginUserDetailsExportAPIView(APIView):
    """Export login user details to Excel or Word format"""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from docx import Document
        from django.http import FileResponse
        from django.utils import timezone
        import io

        export_format = request.query_params.get('export_format', 'excel').lower()
        if export_format not in ['excel', 'word']:
            return Response({'error': 'Invalid export_format. Use "excel" or "word"'}, status=400)

        queryset = LoginUserDetails.objects.all().order_by('-login_time')

        # Filter by user if specified
        user_id = request.query_params.get('user_id')
        if user_id:
            queryset = queryset.filter(user_id=user_id)
        
        # Filter by login status (case-insensitive)
        login_status = request.query_params.get('login_status')
        if login_status:
            queryset = queryset.filter(login_status__iexact=login_status)
        
        # Filter by device type (case-insensitive)
        device_type = request.query_params.get('device_type')
        if device_type:
            queryset = queryset.filter(device_type__iexact=device_type)
        
        # Filter by suspicious flag
        is_suspicious = request.query_params.get('is_suspicious')
        if is_suspicious:
            queryset = queryset.filter(is_suspicious=is_suspicious.lower() == 'true')

        # Filter by year if specified
        year = request.query_params.get('year')
        if year:
            queryset = queryset.filter(login_time__year=year)
            
        # Filter by start date if specified
        start_date = request.query_params.get('start_date')
        if start_date:
            queryset = queryset.filter(login_time__date__gte=start_date)
            
        # Filter by end date if specified
        end_date = request.query_params.get('end_date')
        if end_date:
            queryset = queryset.filter(login_time__date__lte=end_date)

        # Apply search if specified
        search = request.query_params.get('search')
        if search:
            from django.db.models import Q
            queryset = queryset.filter(
                Q(user__username__icontains=search) |
                Q(ip_address__icontains=search) |
                Q(device_name__icontains=search) |
                Q(browser_name__icontains=search)
            )

        if export_format == 'excel':
            return self._export_to_excel(queryset)
        else:
            return self._export_to_word(queryset)

    def _export_to_excel(self, queryset):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from django.http import FileResponse
        from django.utils import timezone
        import io

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Login Details"

        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        headers = ['#', 'Username', 'Login Time', 'Logout Time', 'IP Address', 'Device Type', 'Device Name', 'Browser', 'OS', 'Status', 'Suspicious']
        worksheet.append(headers)

        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        for idx, record in enumerate(queryset, 1):
            login_time_str = timezone.localtime(record.login_time).strftime('%Y-%m-%d %H:%M:%S') if record.login_time else '—'
            logout_time_str = timezone.localtime(record.logout_time).strftime('%Y-%m-%d %H:%M:%S') if record.logout_time else '—'
            
            row = [
                idx,
                record.user.username if record.user else '—',
                login_time_str,
                logout_time_str,
                record.ip_address or '—',
                record.device_type or '—',
                record.device_name or '—',
                f"{record.browser_name or ''} {record.browser_version or ''}".strip() or '—',
                f"{record.os_name or ''} {record.os_version or ''}".strip() or '—',
                record.login_status or '—',
                'Yes' if record.is_suspicious else 'No'
            ]
            worksheet.append(row)
            for cell in worksheet[worksheet.max_row]:
                cell.border = thin_border

        buffer = io.BytesIO()
        workbook.save(buffer)
        buffer.seek(0)

        response = FileResponse(
            buffer,
            as_attachment=True,
            filename=f"login_details_{timezone.now().strftime('%Y%m%d')}.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return response

    def _export_to_word(self, queryset):
        from docx import Document
        from django.http import FileResponse
        from django.utils import timezone
        import io

        doc = Document()
        doc.add_heading('User Login Details Report', 0)

        table = doc.add_table(rows=1, cols=11)
        hdr_cells = table.rows[0].cells
        headers = ['#', 'Username', 'Login Time', 'Logout Time', 'IP Address', 'Device Type', 'Device Name', 'Browser', 'OS', 'Status', 'Suspicious']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for idx, record in enumerate(queryset, 1):
            row_cells = table.add_row().cells
            login_time_str = timezone.localtime(record.login_time).strftime('%Y-%m-%d %H:%M:%S') if record.login_time else '—'
            logout_time_str = timezone.localtime(record.logout_time).strftime('%Y-%m-%d %H:%M:%S') if record.logout_time else '—'
            
            row_cells[0].text = str(idx)
            row_cells[1].text = record.user.username if record.user else '—'
            row_cells[2].text = login_time_str
            row_cells[3].text = logout_time_str
            row_cells[4].text = record.ip_address or '—'
            row_cells[5].text = record.device_type or '—'
            row_cells[6].text = record.device_name or '—'
            row_cells[7].text = f"{record.browser_name or ''} {record.browser_version or ''}".strip() or '—'
            row_cells[8].text = f"{record.os_name or ''} {record.os_version or ''}".strip() or '—'
            row_cells[9].text = record.login_status or '—'
            row_cells[10].text = 'Yes' if record.is_suspicious else 'No'

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = FileResponse(
            buffer,
            as_attachment=True,
            filename=f"login_details_{timezone.now().strftime('%Y%m%d')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return response
